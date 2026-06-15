"""Regression tests — Leg -1 end-to-end round-trip (no network, no JAR).

Complements the unit tests in ``test_legminus1.py`` by exercising the whole
loop against a self-generated DOCX fixture:

    suggest → (simulate human editing the review) → apply → resolved.docx
            → Leg 0 consumes the path-map

Fast and self-contained: the fixture is built in a temp dir with python-docx and
torn down automatically. Skipped if python-docx is unavailable.
"""

from __future__ import annotations

import re
import tempfile
import unittest
from pathlib import Path

import yaml

try:
    import docx  # noqa: F401
    _HAVE_DOCX = True
except ImportError:  # pragma: no cover
    _HAVE_DOCX = False

from velocity_converter.agent_tools import build_velocity_lookup
from velocity_converter.leg0_ingest import apply_path_map
from velocity_converter.legminus1_resolve_paths import run_apply, run_suggest

_REGISTRY = Path(__file__).resolve().parents[2] / "registry" / "path-registry.yaml"


def _build_fixture(path: Path) -> None:
    """A doc with clean leaves, a loop, a loop-only ambiguous leaf, and a miss."""
    from docx import Document
    d = Document()
    d.add_paragraph("Dear {firstName} {lastName},")
    d.add_paragraph("Policy {policyNumber}. Contact {email}.")
    d.add_paragraph("Reference: {bogusLeaf}.")
    d.add_paragraph("[Item]")
    d.add_paragraph("Type {itemTypeCode}, price {purchasePrice}, cover {premium}.")
    d.add_paragraph("[/Item]")
    d.save(str(path))


def _fill_review(review: Path, fills: dict[str, str]) -> None:
    """Simulate a human editing `Final:` lines, block by block."""
    blocks = re.split(r"(?m)^(##\s+Field:.*)$", review.read_text(encoding="utf-8"))
    # re.split keeps the captured headers; rebuild, patching the block after each.
    out, i = [blocks[0]], 1
    while i < len(blocks):
        header, body = blocks[i], blocks[i + 1]
        m = re.match(r"##\s+Field:\s*\{[$+*]?([A-Za-z_][\w.]*)\}", header)
        leaf = m.group(1) if m else ""
        if leaf in fills:
            body = re.sub(r"(?m)^Final:.*$", f"Final: {fills[leaf]}", body)
        out += [header, body]
        i += 2
    review.write_text("".join(out), encoding="utf-8")


@unittest.skipUnless(_HAVE_DOCX, "python-docx not installed")
class TestLegMinus1RoundTrip(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)
        self.stem = "RoundTrip(segment)"
        self.docx = self.root / f"{self.stem}.docx"
        _build_fixture(self.docx)
        self.outdir = self.root / "out"

    def tearDown(self):
        self._tmp.cleanup()

    def _suggest(self) -> Path:
        rc = run_suggest(self.docx, _REGISTRY, self.outdir)
        self.assertEqual(rc, 0)
        d = self.outdir / self.stem
        for name in (".path-review.md", ".path-map.yaml", ".path-changes.md"):
            self.assertTrue((d / f"{self.stem}{name}").exists(), name)
        return d

    def test_suggest_classifies_each_leaf(self):
        d = self._suggest()
        pmap = yaml.safe_load((d / f"{self.stem}.path-map.yaml").read_text())
        by_leaf = {f["leaf"]: f for f in pmap["fields"]}
        # clean exact resolutions
        self.assertEqual(by_leaf["firstName"]["chosen"], "account.data.firstName")
        self.assertEqual(by_leaf["policyNumber"]["chosen"], "policy.policyNumber")
        # loop-scoped resolution
        self.assertEqual(by_leaf["purchasePrice"]["chosen"], "item.data.purchasePrice")
        self.assertEqual(by_leaf["purchasePrice"]["scope"], "loop: Item")
        # loop-only {premium} → within-scope ambiguity (no auto-pick)
        self.assertEqual(by_leaf["premium"]["status"], "ambiguous")
        self.assertEqual(by_leaf["premium"]["chosen"], "")
        # unmatched leaf
        self.assertEqual(by_leaf["bogusLeaf"]["status"], "unmatched")

    def test_apply_produces_resolved_docx_with_provenance(self):
        d = self._suggest()
        review = d / f"{self.stem}.path-review.md"
        # Human resolves the ambiguous + unmatched leaves.
        _fill_review(review, {
            "premium": "item.AccidentalDamage.charges.premium.amount",
            "bogusLeaf": "account.data.email",
        })
        rc = run_apply(review, d)
        self.assertEqual(rc, 0)

        resolved = d / f"{self.stem}.resolved.docx"
        self.assertTrue(resolved.exists())
        from docx import Document
        text = "\n".join(p.text for p in Document(str(resolved)).paragraphs)
        self.assertIn("{account.data.firstName}", text)
        self.assertIn("{item.data.purchasePrice}", text)
        self.assertIn("{item.AccidentalDamage.charges.premium.amount}", text)
        self.assertIn("{account.data.email}", text)
        self.assertIn("[Item]", text)  # loop markers preserved
        self.assertIn("[/Item]", text)

        # Audit records suggested-vs-override provenance.
        audit = (d / f"{self.stem}.path-changes.md").read_text()
        self.assertIn("legMinus1 (suggested)", audit)  # auto-resolved leaves
        self.assertIn("human (selection)", audit)       # filled ambiguous/unmatched

    def test_leg0_consumes_resulting_map(self):
        d = self._suggest()
        review = d / f"{self.stem}.path-review.md"
        _fill_review(review, {
            "premium": "item.AccidentalDamage.charges.premium.amount",
            "bogusLeaf": "account.data.email",
        })
        run_apply(review, d)
        pmap = d / f"{self.stem}.path-map.yaml"

        html = "<p>Dear {firstName}. Items {+itemTypeCode}.</p>"
        out = apply_path_map(html, pmap)
        self.assertIn("{account.data.firstName}", out)
        self.assertIn("{+item.data.itemTypeCode}", out)  # occurrence symbol kept

        # Every non-charge accessor the map emits is a real Leg 0 lookup key.
        lookup = build_velocity_lookup(str(_REGISTRY))
        for f in yaml.safe_load(pmap.read_text())["fields"]:
            acc = f["chosen"]
            if acc and "charges." not in acc:
                self.assertIn(acc, lookup, f"{acc} not resolvable by Leg 0 lookup")


if __name__ == "__main__":
    unittest.main()
