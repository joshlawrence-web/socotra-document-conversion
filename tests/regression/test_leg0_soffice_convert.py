"""Regression tests for the LibreOffice (soffice) docx→HTML converter.

The soffice path is Leg 0's default converter: it preserves the document's
real styling, at the cost of a hard external dependency. These tests guard
the four-stage contract:

  Stage A — _normalize_docx_runs: no {field} / [[$token]] / [Name/] token
            crosses a Word-run boundary (pure python-docx, no soffice needed);
  Stage B — XHTML Writer conversion produces styled, parseable HTML;
  Stage C — _prepare_soffice_html: XHTML cleanup (fonts left as LibreOffice emitted them);
  Stage D — _verify_token_integrity: every token visible in the text exists
            verbatim in the HTML string (the downstream annotators' contract);
  parity  — the field/conditional/loop sets extracted from soffice output
            match the legacy converter's for every pipeline fixture.

Stage B/D/parity tests skip when LibreOffice is not installed.
"""
from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from velocity_converter.leg0_ingest import (
    _FIELD_RE,
    _LOOP_MARKER_RE,
    _VARIANT_MARKER_RE,
    _find_soffice_binary,
    _find_top_level_brackets,
    _normalize_docx_runs,
    _prepare_soffice_html,
    _verify_token_integrity,
    convert_docx,
    convert_docx_soffice,
)

REPO_ROOT = Path(__file__).resolve().parents[2]
FIXTURES = sorted((REPO_ROOT / "tests" / "pipeline" / "fixtures").glob("*.docx"))

SOFFICE_AVAILABLE = _find_soffice_binary() is not None


def _extract_token_sets(html: str) -> tuple[set, set, set]:
    """(field tokens, variant markers, loop markers) visible in the doc text."""
    from bs4 import BeautifulSoup
    text = BeautifulSoup(html, "html.parser").get_text()
    fields = {m.group(0) for m in _FIELD_RE.finditer(text)}
    variants = {m.group(0) for m in _VARIANT_MARKER_RE.finditer(text)}
    loops = {m.group(0) for m in _LOOP_MARKER_RE.finditer(text)}
    return fields, variants, loops


class TestStageANormalization(unittest.TestCase):
    """Stage A alone (no soffice needed): after _normalize_docx_runs, no
    atomic token spans more than one run in any paragraph."""

    def test_fixture_runs_flattened(self):
        from docx import Document

        for fixture in FIXTURES:
            with tempfile.TemporaryDirectory() as td:
                dst = Path(td) / fixture.name
                _normalize_docx_runs(fixture, dst)
                doc = Document(str(dst))
                paras = list(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paras.extend(cell.paragraphs)
                for para in paras:
                    runs = [r.text for r in para.runs]
                    text = "".join(runs)
                    for rx in (_FIELD_RE, _VARIANT_MARKER_RE, _LOOP_MARKER_RE):
                        for m in rx.finditer(text):
                            self.assertTrue(
                                any(m.group(0) in r for r in runs),
                                f"{fixture.name}: token {m.group(0)!r} still "
                                f"spans runs {runs!r}",
                            )

    def test_verifier_catches_fragmented_token(self):
        with self.assertRaises(RuntimeError):
            _verify_token_integrity(
                "<p><span>{first</span><span>Name}</span></p>", "synthetic"
            )

    def test_verifier_catches_empty_body(self):
        with self.assertRaises(RuntimeError):
            _verify_token_integrity("<html><body></body></html>", "synthetic")

    def test_verifier_passes_clean_html(self):
        _verify_token_integrity(
            "<p>Dear {firstName}, [[$stateClause]] [Item/]{price}[/Item]</p>",
            "synthetic",
        )


class TestPrepareSofficeHtml(unittest.TestCase):
    """Stage C — XHTML cleanup; source font names stay untouched."""

    def test_prepare_keeps_source_fonts_and_drops_td_12pt(self):
        raw = """<?xml version="1.0"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
<meta content="application/xhtml+xml; charset=utf-8"/>
<style>
    td, th { vertical-align:top; font-size:12pt;}
    .paragraph-Standard{ font-size:11pt; font-family:Cambria; margin-bottom:0.353cm; }
    .paragraph-Heading_20_1{ font-size:14pt; font-family:Calibri; color:#365f91; font-weight:bold; text-align:left ! important; }
    .cell-Table1_A1{ padding-left:0.191cm; padding-right:0.191cm; padding-top:0cm; padding-bottom:0cm; }
</style>
</head>
<body><p class="paragraph-Standard">Dear {firstName}</p></body>
</html>"""
        html = _prepare_soffice_html(raw)
        self.assertTrue(html.lstrip().startswith("<!DOCTYPE html>"))
        self.assertNotIn("xmlns=", html)
        self.assertIn("text/html; charset=utf-8", html)
        # Source faces must survive — never substitute Helvetica/Times.
        self.assertIn("font-family:Cambria", html)
        self.assertIn("font-family:Calibri", html)
        self.assertNotIn("Helvetica", html)
        self.assertIn("!important", html)
        self.assertNotIn("! important", html)
        # Cell 12pt override removed so Standard 11pt wins inside tables.
        self.assertNotIn("font-size:12pt", html)
        self.assertIn("font-size:11pt", html)
        self.assertIn("{firstName}", html)
        # Default LO cell margins zeroed so table text shares the body edge.
        self.assertIn("padding-left: 0;", html)
        self.assertIn("padding-right: 0;", html)
        self.assertNotIn("padding-left:0.191cm", html)
        self.assertIn("padding-top:0cm", html)  # vertical padding kept


@unittest.skipUnless(SOFFICE_AVAILABLE, "LibreOffice (soffice) not installed")
class TestSofficeConversion(unittest.TestCase):
    def test_styled_output(self):
        """Conversion keeps styling (a <style> block) and real <table> markup."""
        # TestItemsSchedule carries a table + loop markers.
        fixture = next(f for f in FIXTURES if "ItemsSchedule" in f.name)
        html = convert_docx_soffice(fixture)
        self.assertIn("<style", html.lower())
        self.assertIn("<table", html.lower())
        self.assertGreater(len(html), 500)

    def test_body_font_and_size_preserved(self):
        """XHTML filter keeps Normal body size and the document's theme faces.

        StarWriter HTML used to omit body font-family/size entirely. Source
        faces (Calibri/Cambria/…) must remain — register them on the tenant.
        """
        fixture = next(f for f in FIXTURES if "ItemsSchedule" in f.name)
        html = convert_docx_soffice(fixture)
        self.assertRegex(html, r"font-size\s*:\s*11pt")
        self.assertRegex(
            html,
            r"font-family\s*:\s*(Calibri|Cambria|Arial|Times|Helvetica)",
            "theme font-family from the docx must appear in the style block",
        )

    def test_token_integrity_all_fixtures(self):
        for fixture in FIXTURES:
            html = convert_docx_soffice(fixture)  # raises on integrity failure
            self.assertTrue(html.strip())

    def test_token_parity_with_legacy(self):
        """soffice output yields the same field/conditional/loop token sets as
        the legacy converter for every pipeline fixture."""
        for fixture in FIXTURES:
            legacy_sets = _extract_token_sets(convert_docx(fixture))
            soffice_sets = _extract_token_sets(convert_docx_soffice(fixture))
            self.assertEqual(
                legacy_sets, soffice_sets,
                f"token sets diverge for {fixture.name}",
            )
            # Same number of top-level [[...]] blocks in the raw HTML string
            # (annotate_conditionals' position-matching contract).
            self.assertEqual(
                len(_find_top_level_brackets(convert_docx(fixture))),
                len(_find_top_level_brackets(convert_docx_soffice(fixture))),
                f"[[...]] block count diverges for {fixture.name}",
            )


if __name__ == "__main__":
    unittest.main()
