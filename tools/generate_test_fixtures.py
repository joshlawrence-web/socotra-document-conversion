#!/usr/bin/env python3
"""
Generate the canonical test DOCX fixtures for the pipeline test suite.

Usage:
    python3 scripts/generate_test_fixtures.py [--out-dir tests/pipeline/fixtures]

Each document uses real registry accessor paths in {field} tokens and
[[$token]] conditional blocks so Leg 0 + Leg 2 produce high-confidence matches.

Documents:
    TestQuoteSummary(quote).docx      — quote-level fields, 2 conditionals
    TestItemCert(segment).docx        — item/coverage fields, 3 conditionals
    TestRenewalNotice(segment).docx   — policy renewal fields, 3 conditionals
    TestItemsSchedule(segment).docx   — [Item/] loop over the items array, 1 conditional
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path


def _require_docx():
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        return Document, Pt, WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.exit(
            "Missing dependency python-docx.\n"
            "Install with: pip install python-docx --break-system-packages"
        )


def _heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _para(doc, text: str, bold: bool = False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def _table_row(table, label: str, value_token: str):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value_token


# ---------------------------------------------------------------------------
# Document 1: TestQuoteSummary(quote)
# ---------------------------------------------------------------------------

def _build_quote_summary(doc):
    _heading(doc, "ZenCover Quote Summary")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "Thank you for requesting a quote from ZenCover. "
               "Please review the details below.")
    _para(doc, "")

    _heading(doc, "Quote Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Quote Reference", "{quote.quoteNumber}")
    _table_row(tbl, "Cover Start Date", "{quote.startTime}")
    _table_row(tbl, "Cover End Date", "{quote.endTime}")
    _para(doc, "")

    _heading(doc, "Coverage", level=2)
    _para(doc, "Breakdown Cover is included as standard in your plan.")
    _para(doc, "")
    _para(doc, "[[$accidentalDamageNote]]")
    _para(doc, "")
    _para(doc, "[[$theftCoverNote]]")
    _para(doc, "")

    _heading(doc, "Premium Summary", level=2)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.rows[0].cells[0].text = "Charge"
    tbl2.rows[0].cells[1].text = "Amount"
    _table_row(tbl2, "Total Monthly Premium", "{pricing.premiumTotal}")
    _table_row(tbl2, "Other Charges", "{pricing.otherTotal}")
    _para(doc, "")

    _para(doc, "This quote is valid for 30 days from the date of issue.")
    _para(doc, "")
    _para(doc, "ZenCover Customer Services")


# ---------------------------------------------------------------------------
# Document 2: TestItemCert(segment)
# ---------------------------------------------------------------------------

def _build_item_cert(doc):
    _heading(doc, "Item Protection Certificate")

    _para(doc, "This certificate confirms the coverage in force for the item detailed below.")
    _para(doc, "")

    _heading(doc, "Policy Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Policy Number", "{policy.policyNumber}")
    _table_row(tbl, "Certificate Holder", "{account.data.firstName} {account.data.lastName}")
    _table_row(tbl, "Email", "{account.data.email}")
    _table_row(tbl, "Phone", "{account.data.primaryPhone}")
    _para(doc, "")

    _heading(doc, "Covered Item", level=2)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.rows[0].cells[0].text = "Field"
    tbl2.rows[0].cells[1].text = "Value"
    _table_row(tbl2, "Item Type", "{item.data.itemTypeCode}")
    _table_row(tbl2, "Purchase Date", "{item.data.purchaseDate}")
    _table_row(tbl2, "Purchase Price", "{item.data.purchasePrice}")
    _table_row(tbl2, "Serial Number", "{item.data.serialNumber}")
    _para(doc, "")

    _heading(doc, "Item Status at Inception", level=2)
    _para(doc, "[[$warrantyStatusNote]]")
    _para(doc, "")
    _para(doc, "[[$workingOrderNote]]")
    _para(doc, "")

    _heading(doc, "Coverage Summary", level=2)
    _para(doc, "Breakdown Cover is included as standard. Labour and parts are covered.")
    _para(doc, "")
    _para(doc, "[[$accidentalDamageNote]]")
    _para(doc, "")

    _para(doc, "ZenCover Limited")


# ---------------------------------------------------------------------------
# Document 3: TestRenewalNotice(segment)
# ---------------------------------------------------------------------------

def _build_renewal_notice(doc):
    _heading(doc, "Policy Renewal Notice")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "Your ZenCover policy is due for renewal. "
               "Please review the details below.")
    _para(doc, "")

    _heading(doc, "Renewal Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Policy Number", "{policy.policyNumber}")
    _table_row(tbl, "Current Policy End Date", "{policy.data.contractTermEndDate}")
    _table_row(tbl, "Renewal Date", "{policy.data.expectedRenewalDate}")
    _para(doc, "")

    _heading(doc, "Renewal Charges", level=2)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.rows[0].cells[0].text = "Charge"
    tbl2.rows[0].cells[1].text = "Amount"
    _table_row(tbl2, "Total Renewal Amount", "{data.termCharges}")
    _table_row(tbl2, "Settlement Period", "{policy.data.settlementPeriod}")
    _para(doc, "")

    _heading(doc, "Discounts and Adjustments", level=2)
    _para(doc, "[[$loyaltyDiscountNote]]")
    _para(doc, "")

    _heading(doc, "Your Rights", level=2)
    _para(doc, "[[$coolingOffNote]]")
    _para(doc, "")
    _para(doc, "[[$gracePeriodNote]]")
    _para(doc, "")

    _para(doc, "Thank you for choosing ZenCover.")
    _para(doc, "")
    _para(doc, "ZenCover Customer Services")


# ---------------------------------------------------------------------------
# Document 4: TestItemsSchedule(segment) — [Item/] loop over the items array
# ---------------------------------------------------------------------------

def _build_items_schedule(doc):
    _heading(doc, "Covered Items Schedule")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "The items covered under your ZenCover policy are listed below.")
    _para(doc, "")

    _heading(doc, "Your Covered Items", level=2)
    # [Item/]/[/Item] rows mark the repeating section — Leg 0 turns them into
    # a #foreach scaffold so each item renders one row; the header stays once.
    tbl = doc.add_table(rows=1, cols=4)
    tbl.rows[0].cells[0].text = "Item Type"
    tbl.rows[0].cells[1].text = "Purchase Date"
    tbl.rows[0].cells[2].text = "Purchase Price"
    tbl.rows[0].cells[3].text = "Serial Number"
    tbl.add_row().cells[0].text = "[Item/]"
    row = tbl.add_row()
    row.cells[0].text = "{item.data.itemTypeCode}"
    row.cells[1].text = "{item.data.purchaseDate}"
    row.cells[2].text = "{item.data.purchasePrice}"
    row.cells[3].text = "{item.data.serialNumber}"
    tbl.add_row().cells[0].text = "[/Item]"
    _para(doc, "")

    _heading(doc, "Coverage Notes", level=2)
    _para(doc, "Breakdown Cover is included as standard for every listed item.")
    _para(doc, "")
    _para(doc, "[[$accidentalDamageNote]]")
    _para(doc, "")

    _para(doc, "ZenCover Limited")


# ---------------------------------------------------------------------------
# Document 5: TestGiftSchedule(segment) — conditional loop section via the
# loop's when row
# ---------------------------------------------------------------------------

def _build_gift_schedule(doc):
    _heading(doc, "Promotional Gift Schedule")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")

    # The loop itself carries its own conditionality: the customer's seed fills
    # the Item row's `when` in the variants.csv, so Leg 0/2/3 wrap the whole
    # #foreach in #if($data.Item) — no wrapping [[...]] block needed any more.
    _para(doc, "The following promotional gift items are included with your policy:")
    _para(doc, "[Item/]")
    _para(doc, "Gift: {item.data.itemTypeCode} valued at {item.data.purchasePrice}")
    _para(doc, "[/Item]")
    _para(doc, "Gift items are subject to availability.")
    _para(doc, "")

    _para(doc, "[[$theftCoverNote]]")
    _para(doc, "")
    _para(doc, "ZenCover Limited")


# ---------------------------------------------------------------------------
# Document 6: TestStateDisclosure(segment) — N-way variant block ([[$token]])
# ---------------------------------------------------------------------------

def _build_state_disclosure(doc):
    """Exercises the multi-variant conditional feature (the 50-state pattern).

    A single [[$disclosureClause]] variant block selects one of several texts by
    the policy's discountType. Its rows are seeded from condition_seeds.yaml (the
    harness builds the filled variants.csv); the variant texts embed a policy
    *system* field ({policy.policyNumber}) and a policy *custom* field
    ({policy.data.discountAmount}) — both supported wiring categories — so the
    generated if/else-if chain concatenates real accessors and compiles against
    customer-config.jar.
    """
    _heading(doc, "Policy Disclosure Notice")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "Please review the disclosure that applies to your policy below.")
    _para(doc, "")

    _heading(doc, "Policy Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Policy Number", "{policy.policyNumber}")
    _table_row(tbl, "Discount Type", "{policy.data.discountType}")
    _para(doc, "")

    _heading(doc, "Applicable Disclosure", level=2)
    # The variant block: one bare $token, filled from the companion variants.csv.
    _para(doc, "[[$disclosureClause]]")
    _para(doc, "")

    _para(doc, "Thank you for choosing ZenCover.")
    _para(doc, "")
    _para(doc, "ZenCover Customer Services")


# ---------------------------------------------------------------------------
# Document 7: TestVariantThenBinary(segment) — regression for the variant-then-
# binary conditional-form parse bug (fixed 2026-06-16)
# ---------------------------------------------------------------------------

def _build_variant_then_binary(doc):
    """Regression fixture: a [[$token]] variant block placed BEFORE a second,
    plain tokenised block.

    The old conditional-form parser's binary-block regex used a DOTALL non-greedy
    body capture that ran *past* the variant block (which carries no Condition:
    line) and stole the following block's condition — producing two `id: 1`
    entries and dropping the second block. Order is load-bearing: the first
    variant MUST precede the second block to reproduce it. Both blocks are now
    named [[$token]] blocks (bare [[text]] is a hard error) — the suite still
    fails if that parse ever regresses.
    """
    _heading(doc, "Variant Then Binary Notice")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "Please review the disclosure that applies to your policy below.")
    _para(doc, "")

    _heading(doc, "Policy Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Policy Number", "{policy.policyNumber}")
    _table_row(tbl, "Discount Type", "{policy.data.discountType}")
    _para(doc, "")

    _heading(doc, "Applicable Disclosure", level=2)
    # Variant block FIRST (Block 1) …
    _para(doc, "[[$disclosureClause]]")
    _para(doc, "")
    # … second tokenised block SECOND (Block 2) — its condition was the one the
    # old parser stole into the phantom block.
    _para(doc, "[[$coolingOffNote]]")
    _para(doc, "")

    _para(doc, "Thank you for choosing ZenCover.")
    _para(doc, "")
    _para(doc, "ZenCover Customer Services")


# ---------------------------------------------------------------------------
# Document 8: TestVariantBareLeaf(segment) — Decision B (bare leaf in variant text)
# ---------------------------------------------------------------------------

def _build_variant_bare_leaf(doc):
    """Regression fixture for Decision B (variants-only plan §2.4/§2.6).

    A single [[$bareLeafClause]] variant block whose seed variant text embeds a
    **bare leaf** ({discountAmount}) rather than the full accessor. The bare leaf
    was never seen by Leg -1's pass-1 scan (the source doc holds only the
    [[$token]] marker), so Leg 4's variant-text resolver must resolve it to a
    single registry accessor (policy.data.discountAmount) and wire a real Java
    accessor — not silently degrade it to a // TODO. Order/shape mirrors
    TestStateDisclosure; only the seed CSV's leaf form differs.
    """
    _heading(doc, "Bare-Leaf Variant Notice")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "Please review the disclosure that applies to your policy below.")
    _para(doc, "")

    _heading(doc, "Policy Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Policy Number", "{policy.policyNumber}")
    _table_row(tbl, "Discount Type", "{policy.data.discountType}")
    _para(doc, "")

    _heading(doc, "Applicable Disclosure", level=2)
    _para(doc, "[[$bareLeafClause]]")
    _para(doc, "")

    _para(doc, "Thank you for choosing ZenCover.")
    _para(doc, "")
    _para(doc, "ZenCover Customer Services")


# ---------------------------------------------------------------------------
# Document 9: TestNestedVariantLabel(segment) — a [[$x]] reference nested INSIDE
# another variant block's text cell (nested-only label, no document marker)
# ---------------------------------------------------------------------------

def _build_nested_variant_label(doc):
    """Exercises nested [[$x]] references inside a variant text cell.

    The document carries a single [[$benefitClause]] marker. Its seed variant text
    embeds [[$amountLabel]] — a *second* placeholder that has no document marker of
    its own (it exists only as a nested reference). The parse step must synthesize a
    block for amountLabel, condition_dsl peels [[$amountLabel]] → $doc.amountLabel,
    and Leg 4 composes the label's value into benefitClause's plugin string
    (`" + amountLabel + "`), ordering amountLabel's local first via the topo sort.
    Both share policy scope (the nested-label scope-match requirement).
    """
    _heading(doc, "Nested Variant Label Notice")

    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "Please review the benefit that applies to your policy below.")
    _para(doc, "")

    _heading(doc, "Policy Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Policy Number", "{policy.policyNumber}")
    _table_row(tbl, "Discount Type", "{policy.data.discountType}")
    _para(doc, "")

    _heading(doc, "Applicable Benefit", level=2)
    # Only the parent marker is in the document; [[$amountLabel]] lives in its seed
    # text and is synthesized as a nested-only block at parse time.
    _para(doc, "[[$benefitClause]]")
    _para(doc, "")

    _para(doc, "Thank you for choosing ZenCover.")
    _para(doc, "")
    _para(doc, "ZenCover Customer Services")


# ---------------------------------------------------------------------------
# Document 10: TestCoverageGrid(segment) — the "giant table" pattern: one table
# where whole rows appear per coverage. Exercises [Name?] conditional regions
# in all four shapes (in-loop coverage presence / in-loop VALUE condition
# [BreakdownLabourRow?] / doc-level coverage presence / doc-level generic)
# plus coverage-hop dotted fields (always-guarded by Leg 3).
# ---------------------------------------------------------------------------

def _build_coverage_grid(doc):
    _heading(doc, "Your Cover At A Glance")
    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.rows[0].cells[0].text = "Section"
    tbl.rows[0].cells[1].text = "Labour"
    tbl.rows[0].cells[2].text = "Parts"
    tbl.add_row().cells[0].text = "[Item/]"
    row = tbl.add_row()
    row.cells[0].text = "Item: {item.data.itemTypeCode}"
    row.cells[1].text = "{item.data.purchaseDate}"
    row.cells[2].text = "{item.data.purchasePrice}"
    tbl.add_row().cells[0].text = "[AccidentalDamage?]"
    row = tbl.add_row()
    row.cells[0].text = "Accidental Damage"
    row.cells[1].text = "{item.AccidentalDamage.data.labourCovered}"
    row.cells[2].text = "{item.AccidentalDamage.data.partsCovered}"
    tbl.add_row().cells[0].text = "[/AccidentalDamage]"
    row = tbl.add_row()
    row.cells[0].text = "Breakdown"
    row.cells[1].text = "{item.Breakdown.data.labourCovered}"
    row.cells[2].text = "{item.Breakdown.data.partsCovered}"
    tbl.add_row().cells[0].text = "[BreakdownLabourRow?]"
    row = tbl.add_row()
    row.cells[0].text = "Breakdown labour is covered for this item."
    tbl.add_row().cells[0].text = "[/BreakdownLabourRow]"
    tbl.add_row().cells[0].text = "[/Item]"
    tbl.add_row().cells[0].text = "[Theft?]"
    row = tbl.add_row()
    row.cells[0].text = "Theft cover is included on at least one of your items."
    tbl.add_row().cells[0].text = "[/Theft]"
    tbl.add_row().cells[0].text = "[CoolingOffRow?]"
    row = tbl.add_row()
    row.cells[0].text = "You may cancel within your cooling-off period."
    tbl.add_row().cells[0].text = "[/CoolingOffRow]"
    _para(doc, "")
    _para(doc, "ZenCover Limited")


# ---------------------------------------------------------------------------
# Document 11: TestCoverageSchedule(segment) — [Coverage/] plugin-list loop:
# one table row per (item x coverage present), iterated over the plugin-built
# $data.coverages list (registry iterable kind: plugin_list).
# ---------------------------------------------------------------------------

def _build_coverage_schedule(doc):
    _heading(doc, "Coverage Schedule")
    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "The coverages included in your ZenCover policy are listed below.")
    _para(doc, "")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.rows[0].cells[0].text = "Item"
    tbl.rows[0].cells[1].text = "Coverage"
    tbl.rows[0].cells[2].text = "Labour Covered"
    tbl.add_row().cells[0].text = "[Coverage/]"
    row = tbl.add_row()
    row.cells[0].text = "{coverage.itemTypeCode}"
    row.cells[1].text = "{coverage.displayName}"
    row.cells[2].text = "{coverage.labourCovered}"
    tbl.add_row().cells[0].text = "[/Coverage]"
    _para(doc, "")
    _para(doc, "ZenCover Limited")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

FIXTURES = [
    ("TestQuoteSummary(quote).docx", _build_quote_summary),
    ("TestItemCert(segment).docx", _build_item_cert),
    ("TestRenewalNotice(segment).docx", _build_renewal_notice),
    ("TestItemsSchedule(segment).docx", _build_items_schedule),
    ("TestGiftSchedule(segment).docx", _build_gift_schedule),
    ("TestStateDisclosure(segment).docx", _build_state_disclosure),
    ("TestVariantThenBinary(segment).docx", _build_variant_then_binary),
    ("TestVariantBareLeaf(segment).docx", _build_variant_bare_leaf),
    ("TestNestedVariantLabel(segment).docx", _build_nested_variant_label),
    ("TestCoverageGrid(segment).docx", _build_coverage_grid),
    ("TestCoverageSchedule(segment).docx", _build_coverage_schedule),
]


def generate(out_dir: Path) -> list[Path]:
    Document, Pt, _ = _require_docx()
    out_dir.mkdir(parents=True, exist_ok=True)
    written = []
    for filename, builder in FIXTURES:
        doc = Document()
        builder(doc)
        dest = out_dir / filename
        doc.save(str(dest))
        print(f"  wrote {dest}")
        written.append(dest)
    return written


def main():
    parser = argparse.ArgumentParser(description="Generate pipeline test fixture DOCX files.")
    parser.add_argument(
        "--out-dir",
        default="tests/pipeline/fixtures",
        help="Directory to write fixture files (default: tests/pipeline/fixtures)",
    )
    args = parser.parse_args()
    out_dir = Path(args.out_dir)
    print(f"Generating test fixtures → {out_dir}/")
    written = generate(out_dir)
    print(f"\nDone — {len(written)} fixture(s) written.")


if __name__ == "__main__":
    main()
