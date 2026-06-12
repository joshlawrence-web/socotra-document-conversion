#!/usr/bin/env python3
"""
Generate the two extra demo DOCX inputs for the multi-ingest scenario.

Usage:
    python3 tools/generate_demo_docs.py [--out-dir samples/input]

Together with the existing ZenCoverTest(quote).docx and
ZenCoverPolicySummary(segment).docx these give the 4-document demo set
(2 quote-rendered + 2 segment-rendered):

    ZenCoverWelcomeLetter(quote).docx       — account + quote fields, 2 conditionals
    ZenCoverCancellationNotice(segment).docx — policy fields, 2 conditionals
                                               (one with a {field} inside the block,
                                               to show Leg 4 conditional wiring)

Suggested conditions when filling the forms. Use root-prefixed accessor
paths — bare field names (e.g. ``quoteNumber != null``) do NOT translate
to Java accessors and fail the Leg 4 compile check. All of these are
verified to compile against the current ZenCover config:

    WelcomeLetter      1: quote.quoteNumber != null
                       2: quote.startTime != null
    CancellationNotice 1: policy.data.discountAmount != null
                       2: policy.data.coolingOffPeriod != null
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path


def _require_docx():
    try:
        from docx import Document
        return Document
    except ImportError:
        sys.exit(
            "Missing dependency python-docx.\n"
            "Install with: pip install python-docx --break-system-packages"
        )


def _para(doc, text: str):
    return doc.add_paragraph(text)


def _table_row(table, label: str, value_token: str):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value_token


def _build_welcome_letter(doc):
    doc.add_heading("Welcome to ZenCover", level=1)
    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "Welcome aboard! Your ZenCover quote has been prepared — "
               "here is everything you need to know about your new plan.")
    _para(doc, "")

    doc.add_heading("Your Quote", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Quote Reference", "{quoteNumber}")
    _table_row(tbl, "Cover Start Date", "{startTime}")
    _table_row(tbl, "Jurisdiction", "{jurisdiction}")
    _table_row(tbl, "Monthly Premium", "{quotePremiumTotal}")
    _para(doc, "")

    doc.add_heading("Your Benefits", level=2)
    _para(doc, "[[As a new ZenCover customer, a welcome discount has been "
               "applied to your first premium payment.]]")
    _para(doc, "")
    _para(doc, "[[Your plan includes Accidental Damage cover from day one — "
               "no waiting period applies.]]")
    _para(doc, "")

    _para(doc, "We are delighted to have you with us.")
    _para(doc, "")
    _para(doc, "ZenCover Customer Services")


def _build_cancellation_notice(doc):
    doc.add_heading("Policy Cancellation Notice", level=1)
    _para(doc, "Dear {account.data.firstName} {account.data.lastName},")
    _para(doc, "")
    _para(doc, "We confirm that your ZenCover policy has been cancelled "
               "as requested. The details are set out below.")
    _para(doc, "")

    doc.add_heading("Cancellation Details", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    _table_row(tbl, "Policy Number", "{policyNumber}")
    _table_row(tbl, "Cover End Date", "{policy.data.contractTermEndDate}")
    _table_row(tbl, "Settlement Period", "{policy.data.settlementPeriod}")
    _para(doc, "")

    doc.add_heading("Refunds and Charges", level=2)
    _para(doc, "[[A pro-rata refund of {policy.data.discountAmount} will be "
               "issued to your original payment method within 10 working days.]]")
    _para(doc, "")
    _para(doc, "[[Your cooling-off period has not yet expired, so this "
               "cancellation is free of charge.]]")
    _para(doc, "")

    doc.add_heading("What Happens Next", level=2)
    _para(doc, "Your cover will remain in force until the cover end date "
               "shown above. No further payments will be collected.")
    _para(doc, "")
    _para(doc, "ZenCover Limited")


DOCS = [
    ("ZenCoverWelcomeLetter(quote).docx", _build_welcome_letter),
    ("ZenCoverCancellationNotice(segment).docx", _build_cancellation_notice),
]


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate multi-ingest demo DOCX inputs")
    ap.add_argument("--out-dir", default="samples/input")
    args = ap.parse_args()

    Document = _require_docx()
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    for name, builder in DOCS:
        doc = Document()
        builder(doc)
        path = out_dir / name
        doc.save(path)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
