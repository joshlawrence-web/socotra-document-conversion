#!/usr/bin/env python3
"""The demo done-gate: prove a generated .final.vm actually matches its source
Word doc AND has renderingData-correct paths. THIS is the check that stops an
agent declaring success on a template that renders to nothing.

    python3 tools/validate_demo.py "<stem>"            # e.g. ZenCoverProtectionLetter(segment)
    python3 tools/validate_demo.py "<stem>" --registry registry/path-registry.yaml

A "match" means:
  1. No unresolved markers left in the .vm ($TBD_, $doc.cond, {bare}, [Item]).
  2. renderingData shape is correct — every rendering-root-entity field carries
     its entity key. A bare $data.data.<f> or a bare $data.<systemField> renders
     to NOTHING (the plugin never .put()s a "data" key or a flat system field).
     See docs/RenderingDataConfigRelated.md § "rendering-root entity key".
  3. Every static line the author wrote appears in the .vm (.docx source only).
  4. Every [[block]] became a ${data.<key>} ref whose prose is captured in the
     conditional-registry (text lives there, not the .vm — no plugin needed).
     <key> is condN for a binary block, or the token name for a [[$token]]
     variant block (e.g. ${data.disclosureClause}).

Exit 0 = PASS. Exit 1 = MISMATCH (prints every failure). Designed to be re-run
after every fix until it prints PASS — that loop is the whole point.
"""
import argparse
import os
import re
import sys
from pathlib import Path

import yaml

MARKER = re.compile(r"(?<!\$)(?<!\$!)\{[+*$]?([A-Za-z0-9_.]+)\}")   # {leaf}, not ${resolved}/$!{resolved}
BLOCK = re.compile(r"\[\[(.+?)\]\]", re.S)
LOOP = re.compile(r"\[/?[A-Za-z0-9_]+[/?]?\]")  # [Item/] / [Name?] openers, [/Item] closer (+ legacy [Item])
# Entity keys the plugin .put()s — every resolved field must sit under one.
# "coverages" = the plugin-built [Coverage/] list key (the generated plugin puts it).
ENTITY_KEYS = ("quote", "segment", "policy", "account", "pricing", "charges", "termCharges", "coverages")


def _doc_lines(path: Path):
    """Static lines from a .docx (paragraphs + table cells). [] for non-docx."""
    if path.suffix.lower() != ".docx":
        return None
    from docx import Document
    d = Document(str(path))
    out = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            out += [c.text.strip() for c in row.cells if c.text.strip()]
    return out


def _system_field_names(registry: Path | None):
    """Root-entity system field leaf names (system_paths + quote_paths). These
    must never appear bare as $data.<name> — they need the $data.<key> prefix."""
    if not registry or not registry.is_file():
        return set()
    data = yaml.safe_load(registry.read_text()) or {}
    names = set()
    for key in ("system_paths", "quote_paths"):
        for e in data.get(key) or []:
            f = (e.get("field") or "").strip()
            if f:
                names.add(f)
    return names


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("stem")
    ap.add_argument("--registry", default="registry/path-registry.yaml")
    ap.add_argument("--output", default=os.environ.get("CONVERTER_OUTPUT", "workspace/output"))
    ap.add_argument("--inbox", default=os.environ.get("CONVERTER_INBOX", "workspace/inbox"))
    ap.add_argument("--no-config", action="store_true",
                    help="template-only run: config-readiness findings ($TBD_ + "
                         "renderingData shape) become advisory notes, not failures. "
                         "Authoring-faithfulness checks stay hard.")
    args = ap.parse_args()

    stem = args.stem
    out_dir = Path(args.output) / stem
    vm_path = out_dir / f"{stem}.final.vm"
    reg_path = out_dir / f"{stem}.conditional-registry.yaml"
    registry = Path(args.registry)

    if not vm_path.is_file():
        print(f"MISMATCH: no template at {vm_path} — run the pipeline through Leg 3 first.")
        sys.exit(1)
    vm = vm_path.read_text()

    # authoring-gate honesty: a .vm older than its inputs is a stale artifact
    # (a re-fill that never got re-finalized). Hard fail even in no-config.
    stale = []
    vm_mtime = vm_path.stat().st_mtime
    action_dir = Path(os.environ.get("CONVERTER_ACTION", "workspace/action-needed"))
    for inp in (action_dir / f"{stem}.path-review.csv",
                action_dir / f"{stem}.variants.csv",
                out_dir / f"{stem}.mapping.yaml"):
        if inp.is_file() and inp.stat().st_mtime > vm_mtime:
            stale.append(inp.name)
    vm_flat = re.sub(r"\s+", " ", vm)
    # prose split across styling spans (e.g. "[</span>Plan A") still counts
    vm_text = re.sub(r"\s+", " ", re.sub(r"<[^>]+>", "", vm))
    reg = {}
    if reg_path.is_file():
        reg = {b["key"]: b for b in (yaml.safe_load(reg_path.read_text()) or [])}

    doc = None
    for ext in (".docx", ".pdf"):
        p = Path(args.inbox) / f"{stem}{ext}"
        if p.is_file():
            doc = p
            break

    # group A = authoring-faithfulness (always hard). group B = config-readiness
    # (advisory in --no-config: nothing to decide without a product config).
    fails, advisories, checks = [], [], 0

    if stale:
        fails.append(f"stale template: {vm_path.name} is older than its input(s) "
                     f"{sorted(stale)} — re-run finalize before validating")

    # 1. leftover-marker guard
    checks += 1
    if "$TBD_" in vm:
        # config-readiness: no accessor (charge/premium, DataFetcher, un-reviewed
        # marker). Without a config there is nothing to resolve it against.
        tbd = sorted(set(re.findall(r"\$TBD_[A-Za-z0-9_.]+", vm)))
        advisories.append(f"unresolved $TBD_ accessor(s) left in template: {tbd}")
    for bad in ("$doc.cond", "[Item]", "[/Item]"):
        checks += 1
        if bad in vm:
            fails.append(f"unresolved {bad!r} left in template")
    # any surviving loop opener [Name/], region opener [Name?], or closer
    # [/Name] means Leg 0 never consumed the marker (bare [Name] is excluded —
    # it can be prose brackets)
    checks += 1
    leftover = re.findall(r"(?<!\[)\[(?:[A-Za-z_]\w*[/?]|/[A-Za-z_]\w*)\](?!\])", vm)
    if leftover:
        fails.append(f"unconsumed loop/region marker(s) left in template: {sorted(set(leftover))}")
    checks += 1
    if MARKER.search(vm):
        fails.append(f"bare {{leaf}} markers left in template: {MARKER.findall(vm)}")

    # 2. renderingData shape (config-readiness → advisory in --no-config)
    checks += 1
    if "$data.data." in vm:
        advisories.append("renderingData shape: '$data.data.*' has no entity key — "
                          "should be '$data.segment.data.*' (segment) / '$data.quote.data.*' (quote)")
    for sysf in sorted(_system_field_names(registry)):
        checks += 1
        if re.search(r"\$\{?data\." + re.escape(sysf) + r"\b", vm):
            advisories.append(f"renderingData shape: bare '$data.{sysf}' has no entity key — "
                              f"should be '$data.policy.{sysf}' (segment) / '$data.quote.{sysf}' (quote)")
    # any other $data.<x> whose <x> is not a known entity key or a computed/cond key
    for m in re.finditer(r"\$\{?data\.([A-Za-z_]\w*)\.", vm):
        seg = m.group(1)
        checks += 1
        if seg not in ENTITY_KEYS and seg != "data":  # 'data' already reported above
            # tolerate computed top-level keys ($data.disclosureClause.x is unusual;
            # cond refs are $data.condN with no trailing dot, so not matched here)
            advisories.append(f"renderingData shape: '$data.{seg}.*' is not a known entity key "
                              f"({', '.join(ENTITY_KEYS)}) — confirm the plugin .put()s '{seg}'")

    # 3 + 4. doc coverage (docx only)
    if doc and doc.suffix.lower() == ".docx":
        cond_n = 0
        for line in _doc_lines(doc):
            mb = BLOCK.search(line)
            if mb:
                cond_n += 1
                inner = mb.group(1).strip()
                # variant block [[$token]] is keyed by its token name; binary by
                # cond<id> (mirrors leg0_ingest._make_block / _variant_placeholder).
                vtok = re.match(r"^\$([A-Za-z_]\w*)$", inner)
                key = vtok.group(1) if vtok else f"cond{cond_n}"
                src = (reg.get(key) or {}).get("source_text", "")
                checks += 1
                if inner not in src and src not in inner:
                    fails.append(f"[[block]] {key} prose not captured in conditional-registry: {inner!r}")
                checks += 1
                if f"${{data.{key}}}" not in vm and f"$data.{key}" not in vm:
                    fails.append(f"template missing ${{data.{key}}} reference for {key}")
                continue
            if LOOP.fullmatch(line):
                continue
            static = re.sub(r"\s+", " ", LOOP.sub("", MARKER.sub("X", line))).strip()
            for chunk in (c.strip().strip(",.") for c in re.split(r"X|\s{2,}", static)):
                if len(chunk) < 4:
                    continue
                checks += 1
                if chunk not in vm_flat and chunk not in vm_text:
                    fails.append(f"doc prose missing from template: {chunk!r}")
    elif doc:
        print(f"(note: {doc.name} is not .docx — skipped prose-coverage; shape + marker checks only)")
    else:
        print(f"(note: no source doc found in {args.inbox} for '{stem}' — shape + marker checks only)")

    print(f"ran {checks} checks on {vm_path}")

    # with a real config, config-readiness findings are fatal (as before);
    # in --no-config they are advisory only.
    if not args.no_config:
        fails += advisories
        advisories = []

    if fails:
        print(f"\nMISMATCH ({len(fails)}):")
        for f in fails:
            print("  -", f)
        sys.exit(1)
    if advisories:
        print(f"\nNOTE (config-readiness): {len(advisories)} finding(s) — "
              "template-only, nothing to resolve without a product config:")
        for a in advisories:
            print("  -", a)
        print(f"\nPASS (authoring) — template matches the document. "
              f"template-only: {len(advisories)} item(s) need a config before they render.")
        return
    print("PASS — template matches the document and renderingData shape is correct.")


if __name__ == "__main__":
    main()
