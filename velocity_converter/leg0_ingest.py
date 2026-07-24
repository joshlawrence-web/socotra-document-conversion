#!/usr/bin/env python3
"""
Leg 0 — Document Ingestion (PDF / Word → raw HTML) + Field/Conditional Extraction

Converts a customer's source document (PDF or Word) into a rough HTML file
suitable for the existing Leg 1 pipeline, then extracts {field_name} tokens,
named [[$token]] conditional blocks, and [Name/]...[/Name] loop sections,
annotates them, and writes all pipeline-ready artifacts. Bare [[text]] blocks
are a hard error — every conditional must be a named token whose text lives in
the variants.csv. Loop section names must exactly match a registry iterable
name (e.g. [Item/] ... [/Item]); fields inside become the loop's fields in the
mapping and the markers become an #if($doc.<Name>) + #foreach scaffold (the
#if resolves or strips depending on the loop's when-only variants.csv row).

Conditional text ([[$token]] blocks, plus one when-only row per loop) is
handled by a single human-fill file — {stem}.variants.csv — paired with a
machine sidecar {stem}.conditional-blocks.yaml that carries the per-block
metadata the CSV can't (id, render flag, source_text). The legacy
conditional-form.md flow is retained only behind --parse-conditional-form.

Usage:
    python3 -m velocity_converter.leg0_ingest --input <path.docx|path.pdf> [--output-dir <dir>]
    python3 -m velocity_converter.leg0_ingest --input <path.docx|path.pdf> --scan [--output-dir <dir>]
    python3 -m velocity_converter.leg0_ingest --parse-variants-csv <filled.variants.csv> --output-dir <dir>
    python3 -m velocity_converter.leg0_ingest --parse-conditional-form <filled-form.md> --output-dir <dir>  # legacy

Outputs (normal mode):
    {stem}.raw.html                  — raw converted HTML (pre-annotation)
    {stem}.annotated.html            — HTML with {field} → $TBD_field, [[$token]] → $doc.<token>
    {stem}.mapping.yaml              — leg2-compatible mapping (pipeline input; enriched in-place by Leg 2)
    {stem}.conditional-blocks.yaml   — machine sidecar (block metadata) for the parse step
    {stem}.variants.csv              — customer-facing fill file (every conditional block)

Outputs (--scan mode — front-loads the customer handoff):
    {stem}.variants.csv        — customer-facing fill file (only)
    (no machine artifacts — the full ingest writes those, incl. the sidecar, later)

Output (--parse-variants-csv / legacy --parse-conditional-form mode):
    {stem}.conditional-registry.yaml — parsed conditional registry for Leg 4

Exit codes:
    0  success
    1  any error (file not found, unsupported format, import failure)
"""
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

import yaml

from velocity_converter.models import ConditionalRegistry, MappingDoc, block_key, validate_contract
from velocity_converter.workspace import action_needed_file, machine_dir_for_action_file

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _derive_accessor(velocity: str, category: str) -> str:
    """Derive clean accessor from velocity path + category."""
    v = (velocity or "").strip()
    cat = (category or "").strip()
    if cat == "system":
        return "policy." + v[len("$data."):] if v.startswith("$data.") else v
    if cat == "quote_system":
        return "quote." + v[len("$data."):] if v.startswith("$data.") else v
    if cat == "policy_data":
        return "policy." + v[len("$data."):] if v.startswith("$data.") else v
    if v.startswith("$data."):
        return v[len("$data."):]
    if v.startswith("$"):
        return v[1:]
    return v


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


# ---------------------------------------------------------------------------
# Word (.docx) → HTML
# ---------------------------------------------------------------------------

def convert_docx(docx_path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        sys.exit(
            "Missing dependency python-docx.\n"
            "Install with: pip install python-docx --break-system-packages"
        )

    doc = Document(str(docx_path))
    parts: list[str] = []

    def _paragraph_html(para) -> str | None:
        text = para.text.strip()
        if not text:
            return None
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            return f"<h1>{_escape(text)}</h1>"
        if "heading 2" in style_name or "heading 3" in style_name:
            return f"<h2>{_escape(text)}</h2>"
        if "heading" in style_name:
            return f"<h2>{_escape(text)}</h2>"
        runs_with_text = [r for r in para.runs if r.text.strip()]
        if runs_with_text and all(r.bold for r in runs_with_text):
            return f"<p><strong>{_escape(text)}</strong></p>"
        return f"<p>{_escape(text)}</p>"

    def _table_html(table) -> str:
        rows = ["<table>"]
        for row in table.rows:
            cells = "".join(
                f"<td>{_escape(cell.text.strip())}</td>" for cell in row.cells
            )
            rows.append(f"  <tr>{cells}</tr>")
        rows.append("</table>")
        return "\n".join(rows)

    from docx.oxml.ns import qn  # noqa: F401,PLC0415

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph  # noqa: PLC0415
            para = Paragraph(child, doc)
            html = _paragraph_html(para)
            if html:
                parts.append(html)

        elif tag == "tbl":
            from docx.table import Table  # noqa: PLC0415
            table = Table(child, doc)
            parts.append(_table_html(table))

    body_content = "\n".join(parts)
    return f"<html>\n<body>\n{body_content}\n</body>\n</html>\n"


# ---------------------------------------------------------------------------
# Word (.docx) → HTML via LibreOffice headless (lossless styling path)
# ---------------------------------------------------------------------------
# The legacy convert_docx() above keeps only heading level / all-bold / table
# structure; every colour, font, size and alignment in the source document is
# discarded. The soffice path converts with LibreOffice's real layout engine,
# so the document's actual styling (a <style> block + inline CSS resolved from
# the docx style/theme XML) survives into raw.html → annotated.html → final.vm.
#
# Filter choice: ``xhtml:XHTML Writer File`` (not ``html:HTML (StarWriter)``).
# StarWriter HTML drops body font-family/size (only heading .western rules
# survive); the XHTML filter emits per-style classes with the real theme fonts
# and sizes (e.g. ``.paragraph-Standard { font-family:Cambria; font-size:11pt }``).
# Source font names are kept as-is (Calibri/Cambria/…) — register those faces
# on the Socotra tenant; do not substitute PDF built-ins.
#
# Stage C cleans the XHTML for Velocity (DOCTYPE/xmlns, ``!important``, and the
# XHTML ``td { font-size:12pt }`` override that fights the document body size).
#
# The one hazard: soffice emits one inline tag per Word *run*, and Word splits
# runs arbitrarily (spell-check artifacts), so a {field} / [[$token]] /
# [Name/] marker can straddle a run boundary and come out fragmented across
# tags — which silently breaks the downstream annotation regexes that match
# tokens on the raw HTML string. Fix is deterministic and lives in the docx,
# not the HTML: flatten any token-spanning runs before conversion (no run
# boundary inside a token ⇒ soffice cannot split it), then *verify* the HTML
# and fail loudly — never repair markup after the fact.

_SOFFICE_HINT = (
    "LibreOffice (soffice) is required to convert .docx documents.\n"
    "Install with: brew install --cask libreoffice   (macOS)\n"
    "          or: apt install libreoffice            (Debian/Ubuntu)\n"
    "Or re-run with --converter legacy for the style-less fallback converter."
)

# Document-body conditional markers are always a short named token (bare
# [[text]] blocks are rejected by extract_conditionals).
_VARIANT_MARKER_RE = re.compile(r"\[\[\$[A-Za-z_]\w*\]\]")

# LibreOffice XHTML Writer filter — keeps theme fonts/sizes that StarWriter HTML drops.
_SOFFICE_FILTER = "xhtml:XHTML Writer File"


def _find_soffice_binary() -> str | None:
    """Locate the soffice binary (PATH, then the macOS app bundle)."""
    found = shutil.which("soffice")
    if found:
        return found
    mac_path = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    return str(mac_path) if mac_path.exists() else None


def _atomic_token_spans(text: str) -> list[tuple[int, int]]:
    """Character spans in ``text`` that must never cross a Word-run boundary:
    {field} placeholders, [[$token]] markers, and [Name/] / [/Name] loop
    markers (the exact patterns the downstream annotators match on the raw
    HTML string)."""
    spans: list[tuple[int, int]] = []
    for rx in (_FIELD_RE, _VARIANT_MARKER_RE, _LOOP_MARKER_RE):
        spans.extend(m.span() for m in rx.finditer(text))
    return spans


def _flatten_token_runs(paragraph) -> None:
    """Rewrite ``paragraph``'s runs so no atomic token spans more than one run.

    Same technique as Leg -1's placeholder rewrite: the runs a token touches
    are collapsed into the first (full combined text there, the rest blanked).
    Mid-token formatting collapses to the first run's — the honest outcome: a
    placeholder resolves to one value, it cannot render two half-styles anyway.
    """
    runs = paragraph.runs
    if len(runs) < 2:
        return
    text = "".join(r.text for r in runs)
    spans = _atomic_token_spans(text)
    if not spans:
        return
    # Cumulative end offset of each run within the paragraph text.
    ends: list[int] = []
    pos = 0
    for r in runs:
        pos += len(r.text)
        ends.append(pos)

    def run_index(char_pos: int) -> int:
        for i, end in enumerate(ends):
            if char_pos < end:
                return i
        return len(runs) - 1

    # Collapse right-to-left so earlier offsets stay valid in `ends`.
    for start, end in sorted(spans, reverse=True):
        first, last = run_index(start), run_index(end - 1)
        if first == last:
            continue
        combined = "".join(runs[i].text for i in range(first, last + 1))
        runs[first].text = combined
        for i in range(first + 1, last + 1):
            runs[i].text = ""


def _normalize_docx_runs(src: Path, dst: Path) -> None:
    """Write a copy of ``src`` with token-spanning runs flattened (Stage A).

    Also strips ``w:embedTrueTypeFonts`` from settings.xml: LibreOffice honors
    it on XHTML export and base64-embeds local system fonts — one oversized
    font blows libxml's text-node limit and the whole conversion hard-fails
    ("huge text node"). Embedding is a Word save preference, never content.
    """
    from docx import Document  # noqa: PLC0415

    doc = Document(str(src))
    for para in doc.paragraphs:
        _flatten_token_runs(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _flatten_token_runs(para)
    settings = doc.settings.element
    for tag in ("embedTrueTypeFonts", "saveSubsetFonts"):
        for el in settings.findall(
                f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{tag}"):
            settings.remove(el)
    doc.save(str(dst))


def _verify_token_integrity(html: str, source_name: str) -> None:
    """Every token visible in the document text must exist verbatim in the raw
    HTML string (Stage D — verifier only). A failure means the Stage A
    flattening has a gap; fix it there, never patch the HTML."""
    try:
        from bs4 import BeautifulSoup  # noqa: PLC0415
        text = BeautifulSoup(html, "html.parser").get_text()
    except ImportError:
        text = re.sub(r"<[^>]+>", " ", html)

    if not text.strip():
        raise RuntimeError(
            f"soffice produced an empty document for {source_name} — conversion failed"
        )
    missing = sorted({
        m.group(0)
        for rx in (_FIELD_RE, _VARIANT_MARKER_RE, _LOOP_MARKER_RE)
        for m in rx.finditer(text)
        if m.group(0) not in html
    })
    if missing:
        raise RuntimeError(
            f"token(s) fragmented across tags in soffice HTML for {source_name}: "
            + ", ".join(missing)
            + "\nThis is a gap in Leg 0's run-flattening (_flatten_token_runs) — "
            "fix the flattening, do not patch the HTML."
        )


def _prepare_soffice_html(html: str) -> str:
    """Stage C — clean LibreOffice XHTML for Velocity without changing fonts.

    Keeps theme ``font-family`` values exactly as LibreOffice resolved them
    (Calibri, Cambria, …). Also:

    - strips XHTML chrome (DOCTYPE/xmlns)
    - fixes invalid ``! important``
    - drops the XHTML ``td/th { font-size:12pt }`` override that fights the
      document's Normal/Standard body size
    - zeroes horizontal table-cell padding (``padding-left`` / ``padding-right``)
      that the XHTML filter copies from Word's default cell margins (~0.075in).
      That padding shifts every table column right of the body text edge; for
      borderless letter tables the PDF then looks indented vs the .docx.
    """
    # Drop XML prologue / MathML doctype — Velocity templates are HTML.
    html = re.sub(r"^\s*<\?xml[^>]*\?>\s*", "", html, count=1, flags=re.I)
    html = re.sub(
        r"<!DOCTYPE[^>]*>\s*",
        "<!DOCTYPE html>\n",
        html,
        count=1,
        flags=re.I | re.S,
    )
    html = re.sub(r"\sxmlns(?::\w+)?=\"[^\"]*\"", "", html)
    html = re.sub(
        r'content="application/xhtml\+xml;\s*charset=([^"]+)"',
        r'content="text/html; charset=\1"',
        html,
        flags=re.I,
    )
    # LibreOffice emits the invalid ``! important`` (space) form.
    html = html.replace("! important", "!important")

    def _style_block(match: re.Match) -> str:
        css = match.group(1)
        # XHTML filter defaults table cells to 12pt, which overrides the
        # document's Normal/Standard 11pt inside every <td>.
        css = re.sub(
            r"(td\s*,\s*th\s*\{[^}]*?)font-size\s*:\s*[^;]+;?\s*",
            r"\1",
            css,
            flags=re.I,
        )
        # Flush table text to the body margin: drop LO's default cell margins.
        css = re.sub(
            r"padding-(left|right)\s*:\s*[^;]+;",
            r"padding-\1: 0;",
            css,
            flags=re.I,
        )
        return f"<style>{css}</style>"

    return re.sub(
        r"<style[^>]*>(.*?)</style>",
        _style_block,
        html,
        count=1,
        flags=re.I | re.S,
    )


def convert_docx_soffice(docx_path: Path) -> str:
    """Convert .docx → styled HTML via LibreOffice headless (Stages A+B+C+D)."""
    soffice = _find_soffice_binary()
    if soffice is None:
        sys.exit(_SOFFICE_HINT)

    tmp = Path(tempfile.mkdtemp(prefix="leg0-soffice-"))
    try:
        normalized = tmp / docx_path.name
        _normalize_docx_runs(docx_path, normalized)
        outdir = tmp / "out"
        # Unique profile per invocation — soffice can't share one profile
        # across concurrent runs.
        profile_uri = (tmp / "profile").as_uri()
        result = subprocess.run(
            [
                soffice, "--headless", "--norestore",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to", _SOFFICE_FILTER,
                "--outdir", str(outdir), str(normalized),
            ],
            capture_output=True, text=True, timeout=120,
        )
        out_file = next(
            (
                p for p in sorted(outdir.glob(f"{normalized.stem}.*"))
                if p.suffix.lower() in {".xhtml", ".html", ".htm"}
            ),
            None,
        )
        if result.returncode != 0 or out_file is None:
            sys.exit(
                f"soffice conversion failed for {docx_path.name} "
                f"(rc={result.returncode}):\n{result.stderr.strip()}"
            )
        html = out_file.read_text(encoding="utf-8", errors="replace")
        html = _prepare_soffice_html(html)
        _verify_token_integrity(html, docx_path.name)
        return html
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


# ---------------------------------------------------------------------------
# PDF → HTML
# ---------------------------------------------------------------------------

def convert_pdf(pdf_path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        sys.exit(
            "Missing dependency pdfplumber.\n"
            "Install with: pip install pdfplumber --break-system-packages"
        )

    parts: list[str] = []

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            if page_num > 0:
                parts.append("<!-- page break -->")

            table_bboxes: list[tuple] = []
            try:
                tables = page.extract_tables()
                for table_data in (tables or []):
                    if not table_data:
                        continue
                    rows = ["<table>"]
                    for row in table_data:
                        cells = "".join(
                            f"<td>{_escape((cell or '').strip())}</td>"
                            for cell in (row or [])
                        )
                        rows.append(f"  <tr>{cells}</tr>")
                    rows.append("</table>")
                    parts.append("\n".join(rows))
                for t_obj in (page.find_tables() or []):
                    table_bboxes.append(t_obj.bbox)
            except Exception:
                pass

            try:
                chars = page.chars or []
                if chars:
                    avg_height = sum(c.get("height", 0) for c in chars) / len(chars)
                else:
                    avg_height = 0

                def _in_table(c) -> bool:
                    x0, y0, x1, y1 = c.get("x0", 0), c.get("top", 0), c.get("x1", 0), c.get("bottom", 0)
                    for tx0, ty0, tx1, ty1 in table_bboxes:
                        if x0 >= tx0 and x1 <= tx1 and y0 >= ty0 and y1 <= ty1:
                            return True
                    return False

                text_chars = [c for c in chars if not _in_table(c)]

                lines: dict[float, list] = {}
                for c in text_chars:
                    top = round(c.get("top", 0), 1)
                    lines.setdefault(top, []).append(c)

                for top_pos in sorted(lines):
                    line_chars = sorted(lines[top_pos], key=lambda c: c.get("x0", 0))
                    line_text = "".join(c.get("text", "") for c in line_chars).strip()
                    if not line_text:
                        continue
                    line_avg_h = sum(c.get("height", 0) for c in line_chars) / len(line_chars)
                    if avg_height > 0 and line_avg_h > avg_height * 1.2:
                        parts.append(f"<h2>{_escape(line_text)}</h2>")
                    else:
                        parts.append(f"<p>{_escape(line_text)}</p>")

            except Exception:
                raw = page.extract_text() or ""
                for line in raw.splitlines():
                    line = line.strip()
                    if line:
                        parts.append(f"<p>{_escape(line)}</p>")

    body_content = "\n".join(parts)
    return f"<html>\n<body>\n{body_content}\n</body>\n</html>\n"


# ---------------------------------------------------------------------------
# Field extraction (E-T1)
# ---------------------------------------------------------------------------

# Group 1: optional occurrence symbol ($ optional, + one-or-more, * zero-or-more;
# bare = required). Group 2: field name. See models.OCCURRENCE_SYMBOLS.
_FIELD_RE = re.compile(r"\{([$+*]?)([a-zA-Z_][a-zA-Z0-9_.]*)\}")

# Accessor roots that map to the rendering-root entity in renderingData. Their
# registry velocity is root-relative ($data.X / $data.data.X) and needs the
# $data.<root> entity key spliced in. account.* / item.* name their own
# renderingData key and are excluded.
_ROOT_ENTITY_ACCESSORS = frozenset({"policy", "quote", "segment"})


def extract_fields(html: str, registry_path=None, rendering_root=None) -> list[dict]:
    """Extract {field_name} tokens from HTML (on plain text). Deduplicates.

    An occurrence symbol may prefix the name — {$x} optional, {x} required,
    {+x} one or more, {*x} zero or more — and is recorded as a normalized
    ``occurrence`` value on the field. The canonical token never carries the
    symbol ($TBD_x for all four forms). Conflicting symbols for the same name
    keep the first one seen and warn on stderr.

    When registry_path is provided, dotted-path placeholders (e.g. {account.data.firstName})
    are resolved against the registry and their velocity path written into data_source.
    Unresolved dotted names get data_source = "UNRESOLVED:<name>" and a stderr warning.
    """
    from velocity_converter.models import OCCURRENCE_SYMBOLS

    try:
        from bs4 import BeautifulSoup
        text = BeautifulSoup(html, "html.parser").get_text()
    except ImportError:
        text = re.sub(r"<[^>]+>", " ", html)

    seen: dict[str, list[int]] = {}
    occurrences: dict[str, str] = {}
    conflicts: list[str] = []
    for i, m in enumerate(re.finditer(_FIELD_RE, text)):
        symbol, name = m.group(1), m.group(2)
        seen.setdefault(name, []).append(i)
        if name not in occurrences:
            occurrences[name] = OCCURRENCE_SYMBOLS[symbol]
        elif occurrences[name] != OCCURRENCE_SYMBOLS[symbol]:
            conflicts.append(name)

    lookup: dict = {}
    meta_lookup: dict = {}
    if registry_path and any("." in name for name in seen):
        try:
            _scripts = str(Path(__file__).parent)
            if _scripts not in sys.path:
                sys.path.insert(0, _scripts)
            from agent_tools import (  # noqa: PLC0415
                build_velocity_lookup,
                build_velocity_meta_lookup,
                render_root_velocity,
            )
            lookup = build_velocity_lookup(registry_path)
            meta_lookup = build_velocity_meta_lookup(registry_path)
        except Exception:
            pass

    unresolved: list[str] = []
    fields: list[dict] = []
    for name in seen:
        candidate = None
        if "." in name and lookup:
            velocity = lookup.get(name)
            if velocity:
                # Rendering-root-entity fields (policy/quote/segment accessors)
                # render under the doc's named root key in renderingData
                # ($data.segment.* / $data.quote.*). The registry velocity is
                # root-relative, so splice the key in — mirrors Leg 2's per-root
                # verdict (_reprefix) for the JAR path. account.* / item.* name
                # their own renderingData key and are left untouched.
                if name.split(".", 1)[0] in _ROOT_ENTITY_ACCESSORS:
                    velocity = render_root_velocity(velocity, rendering_root)
                data_source = velocity
                # DataFetcher-sourced paths (e.g. account.data.firstName →
                # getAccount) share their velocity path with a phantom direct
                # row, so the path alone can't drive the fetch. Stamp a
                # candidate block here so Leg 4 wires the DataFetcher call.
                meta = meta_lookup.get(name)
                if meta and meta.get("source") == "datafetcher":
                    candidate = {
                        "path": meta.get("velocity") or velocity,
                        "match_step": "exact",
                        "source": "datafetcher",
                        "datafetcher_method": meta.get("datafetcher_method", ""),
                        "datafetcher_arg": meta.get("datafetcher_arg"),
                        "datafetcher_key": meta.get("datafetcher_key", ""),
                    }
                    if meta.get("valid_roots"):
                        candidate["valid_roots"] = meta["valid_roots"]
            else:
                data_source = f"UNRESOLVED:{name}"
                unresolved.append(name)
        else:
            data_source = ""
        field = {
            "name": name,
            "token": f"$TBD_{name}",
            "data_source": data_source,
            "confidence": "",
            "occurrence": occurrences.get(name, "required"),
        }
        if candidate:
            field["candidate"] = candidate
        fields.append(field)

    if conflicts:
        print(
            "WARNING: conflicting occurrence symbols for: "
            f"{', '.join(sorted(set(conflicts)))} — first symbol seen wins",
            file=sys.stderr,
        )
    if unresolved:
        print(
            f"WARNING: unresolved dotted placeholders: {', '.join(unresolved)}",
            file=sys.stderr,
        )

    return fields


def apply_path_map(html: str, path_map_path: Path) -> str:
    """Rewrite bare ``{leaf}`` → ``{accessor}`` using a Leg -1 path-map.

    The map (``<stem>.path-map.yaml``) is the human-validated leaf→accessor
    resolution from Leg -1. Occurrence symbols are preserved. Leaves with an
    empty/absent ``chosen`` are left untouched. Returns the rewritten HTML.
    """
    try:
        data = yaml.safe_load(path_map_path.read_text(encoding="utf-8")) or {}
    except Exception as exc:  # noqa: BLE001
        print(f"WARNING: could not read path-map {path_map_path}: {exc}", file=sys.stderr)
        return html
    mapping = {
        f["leaf"]: f["chosen"]
        for f in (data.get("fields") or [])
        if f.get("leaf") and f.get("chosen")
    }
    if not mapping:
        return html

    applied: set[str] = set()

    def _repl(m: re.Match) -> str:
        symbol, name = m.group(1), m.group(2)
        if name in mapping:
            applied.add(name)
            return "{" + symbol + mapping[name] + "}"
        return m.group(0)

    result = _FIELD_RE.sub(_repl, html)
    print(f"Applied path-map: rewrote {len(applied)}/{len(mapping)} leaf placeholder(s).")
    return result


def annotate_fields(html: str, fields: list[dict]) -> str:
    """Replace {field_name} → $TBD_field_name in the HTML string.

    Occurrence symbols are accepted and stripped: {$x}, {+x}, {*x} all become
    the same canonical token as {x}.
    """
    tokens = {f["name"]: f["token"] for f in fields}
    return _FIELD_RE.sub(lambda m: tokens.get(m.group(2), m.group(0)), html)


# ---------------------------------------------------------------------------
# Conditional block extraction (E-T2)
# ---------------------------------------------------------------------------

def _find_top_level_brackets(text: str) -> list[tuple[int, int, str]]:
    """Find top-level [[...]] blocks, skipping nested ones.
    Returns list of (start, end, inner_content) tuples.
    Handles cases like [[ outer [[inner]] still outer ]] correctly.
    """
    results = []
    i = 0
    n = len(text)
    while i < n - 1:
        if text[i : i + 2] == "[[":
            depth = 1
            start = i
            i += 2
            while i < n - 1 and depth > 0:
                if text[i : i + 2] == "[[":
                    depth += 1
                    i += 2
                elif text[i : i + 2] == "]]":
                    depth -= 1
                    if depth == 0:
                        results.append((start, i + 2, text[start + 2 : i]))
                    i += 2
                else:
                    i += 1
            if depth > 0:
                i = start + 2  # unclosed — skip
        else:
            i += 1
    return results


# A *variant* block tokenises its content with a single bare $identifier:
# [[$stateClause]] (the 50-state feature). The leading $ + single-token (no
# spaces/punctuation) discriminates it from a binary [[literal text]] block.
_VARIANT_TOKEN_RE = re.compile(r"^\$([A-Za-z_]\w*)$")


def _variant_placeholder(source_text: str) -> str | None:
    """Return the placeholder name if ``source_text`` is a variant token.

    ``[[$stateClause]]`` → ``stateClause``; anything else returns None (and is
    rejected by :func:`extract_conditionals` — bare blocks are unsupported).
    """
    s = (source_text or "").strip()
    m = _VARIANT_TOKEN_RE.match(s)
    return m.group(1) if m else None


def extract_conditionals(html: str) -> list[dict]:
    """Extract ``[[$token]]`` variant blocks. Bare blocks are rejected.

    Every conditional in the document must be a named variant token —
    ``[[$stateClause]]`` — whose text/conditions the customer supplies in the
    variants.csv. A bare ``[[literal text]]`` block, a malformed token, or a
    token used twice raises :class:`ValueError` listing every offender (nothing
    is written). Conditional *nesting* is authored in the CSV via ``[[$other]]``
    references inside a ``text`` cell, never in the document body.
    """
    try:
        from bs4 import BeautifulSoup
        text = BeautifulSoup(html, "html.parser").get_text()
    except ImportError:
        text = re.sub(r"<[^>]+>", " ", html)

    blocks: list[dict] = []
    errors: list[str] = []
    seen: set[str] = set()
    for block_id, (_start, _end, content) in enumerate(_find_top_level_brackets(text), start=1):
        source = content.strip()
        ph = _variant_placeholder(source)
        if ph is None:
            snippet = source if len(source) <= 60 else source[:57] + "…"
            errors.append(
                f"[[{snippet}]] — bare [[text]] blocks are not supported; name the "
                f"block [[$myToken]] and put its text in the variants.csv"
            )
            continue
        if ph in seen:
            # A block's text lives in the CSV keyed by name, so a repeated
            # marker is by definition the same content — legitimate reuse
            # (e.g. a shared state-label inside many benefit blocks). One
            # block is registered; the annotator's literal-replace path
            # annotates every occurrence with the same $doc.<token>.
            print(
                f"NOTE: [[${ph}]] appears more than once — every occurrence "
                "renders the same CSV-defined text",
                file=sys.stderr,
            )
            continue
        seen.add(ph)
        blocks.append({
            "id": block_id,
            "key": ph,
            "placeholder": ph,
            "variant": True,
            "source_text": source,
            "raw_text": content,
            "top_level": True,
        })

    if errors:
        # A bare block spanning a <table> gets a targeted hint (tables can't
        # live in a CSV text cell; the stripped-text capture would garble them).
        if any("<table" in c for _s, _e, c in _find_top_level_brackets(html)):
            errors.append(
                "note: a [[...]] block spanning a <table> cannot be tokenised — "
                "tables inside conditional text are not supported; leave the table "
                "outside the conditional markers"
            )
        raise ValueError(
            "conditional block validation failed:\n  - " + "\n  - ".join(errors)
        )
    return blocks


def annotate_conditionals(html: str, blocks: list[dict]) -> str:
    """Replace [[$token]] → [[$token]]$doc.<token> for every variant block.

    Blocks are matched by character position (right-to-left) when the HTML
    bracket count agrees with the block list, else by literal string replace.
    """
    regions = _find_top_level_brackets(html)

    if len(regions) != len(blocks):
        result = html
        for b in blocks:
            original = "[[" + b.get("raw_text", b["source_text"]) + "]]"
            result = result.replace(original, f"{original}$doc.{b['key']}")
    else:
        result = html
        for (start, end, content), b in zip(reversed(regions), reversed(blocks)):
            label = f"[[{content}]]"
            result = result[:start] + f"{label}$doc.{b['key']}" + result[end:]

    return result


# ---------------------------------------------------------------------------
# Loop section extraction ([Name/] ... [/Name])
# ---------------------------------------------------------------------------

# Single-bracket loop markers: [Name/] opens a repeating section, [/Name]
# closes it (the trailing slash marks the opener as a loop, not prose in
# brackets). Lookarounds exclude [[conditional]] double brackets. The name
# must exactly match a registry iterable name (e.g. [Item/]) — Leg 2 resolves
# loop roots by exact name only. Group 1 = leading slash (closer), group 3 =
# trailing slash (loop opener) or `?` (conditional-region opener, e.g.
# [AccidentalDamage?] … [/AccidentalDamage] — wraps the enclosed rows in a
# presence/condition guard instead of a #foreach); a plain [Name] is legacy
# syntax and only draws a migration warning when its [/Name] closer is present.
_LOOP_MARKER_RE = re.compile(r"(?<!\[)\[(/?)([A-Za-z_]\w*)([/?]?)\](?!\])")

# A directive that ended up as the sole content of a paragraph or a table row
# (other cells empty) collapses to a bare line so Leg 3's line-level #foreach
# replacement / unregistered-guard strip applies — and so marker rows don't
# render as blank gaps in the PDF. Openers may carry their #if($doc.<Name>)
# guard on the line above the #foreach (and closers a second #end); a
# conditional-region marker leaves a lone #if($doc.<Name>) / #if($item.<Cov>)
# opener. Allow attributes on <p>/<td>/<tr> — the XHTML Writer filter emits
# class/style on every tag; "empty" sibling cells often hold a nbsp <p>, not a
# bare </td>.
_DIRECTIVE = (
    r"(?:#if\(\$doc\.\w+\)\n)?#foreach \([^)]*\)"
    r"|#if\(\$\w+\.\w+\)"
    r"|#end(?:\n#end)?"
)
_EMPTY_P = (
    r"<p(?:\s[^>]*)?>\s*(?:&nbsp;|&#160;|\u00a0|<br\s*/?>|\s)*</p>"
)
_EMPTY_TD = rf"<td(?:\s[^>]*)?>\s*(?:{_EMPTY_P}\s*)*</td>"
# Directive may still be wrapped in <p class=…> when it sits mid-line inside a
# <td> (paragraph peel only matches a <p> that starts a line).
_DIRECTIVE_CELL = rf"(?:<p(?:\s[^>]*)?>\s*)?({_DIRECTIVE})(?:\s*</p>)?"
_DIRECTIVE_LINE_CLEANUP = [
    (re.compile(rf"^(\s*)<p(?:\s[^>]*)?>({_DIRECTIVE})</p>\s*$", re.M), r"\1\2"),
    (
        # Collapse a marker <tr> to a bare directive on its own line(s). Leading
        # / trailing newlines matter: Leg 3's unregistered-guard strip is
        # line-based (`#if($doc.X)` must be the whole line), and gluing
        # `#foreach` onto the next `<tr>` hides the data row when the guard
        # survives as `#if($data.Item)`.
        re.compile(
            rf"<tr(?:\s[^>]*)?>\s*<td(?:\s[^>]*)?>\s*{_DIRECTIVE_CELL}\s*</td>"
            rf"(?:\s*{_EMPTY_TD})*\s*</tr>",
            re.M,
        ),
        r"\n\1\n",
    ),
]


def _registry_coverage_map(registry_path) -> tuple[dict, dict]:
    """``({coverage_name: {exposure, list_method}}, {iterable_name: iterator})``.

    Used to classify a ``[Name?]`` conditional-region marker: a name matching a
    coverage gets automatic presence wiring (``list_method`` is the exposure
    list's Java accessor, e.g. ``items`` from ``$data.items`` — Leg 4 walks
    ``<root>.<list_method>()`` to test presence); anything else is a generic
    conditional region the customer conditions via variants.csv. The second
    map gives each iterable's registry iterator (e.g. ``Item → item``) for
    in-loop value regions, whose ``when`` paths must root at that iterator.
    """
    if not registry_path:
        return {}, {}
    try:
        reg = yaml.safe_load(Path(registry_path).read_text(encoding="utf-8")) or {}
    except Exception:
        return {}, {}
    iter_methods: dict = {}
    iterators: dict = {}
    for it in reg.get("iterables") or []:
        lv = str(it.get("list_velocity") or "")
        name = str(it.get("name") or "")
        if name and lv.startswith("$data.") and "." not in lv[len("$data."):]:
            iter_methods[name] = lv[len("$data."):]
        if name:
            iterators[name] = str(it.get("iterator") or "").lstrip("$")
    cov_map: dict = {}
    for exp in reg.get("exposures") or []:
        exp_name = str(exp.get("name") or "")
        for cov in exp.get("coverages") or []:
            name = str(cov.get("name") or "")
            if name and exp_name:
                cov_map.setdefault(name, {
                    "exposure": exp_name,
                    "list_method": iter_methods.get(exp_name, ""),
                    # the exposure's registry iterator (e.g. "item") — the
                    # in-loop guard must reference the same variable Leg 3
                    # splices into the #foreach, not a name.lower() guess
                    "iterator": iterators.get(exp_name, ""),
                })
    return cov_map, iterators


def extract_loops(
    html: str,
    fields: list[dict],
    cond_blocks: list[dict] | None = None,
    registry_path=None,
) -> tuple[str, list[dict], list[dict]]:
    """Detect [Name/]...[/Name] repeating sections and [Name?]...[/Name]
    conditional regions in annotated HTML.

    Returns (html, top_level_fields, loops):
      - loop markers become ``#if($doc.<Name>)`` + ``#foreach ($<name> in
        $TBD_<Name>)`` / ``#end`` + ``#end`` scaffold lines (Leg 2 supplies the
        real directive, Leg 3 swaps it in and strips the ``#if`` guard when the
        loop's ``when`` row was left blank);
      - a field whose every occurrence falls inside one loop section moves from
        the top level into that loop's ``fields`` list;
      - loops is a list of {name, token, iterator, fields} dicts.

    Every loop also appends a ``render: template`` block (key = loop name) to
    ``cond_blocks`` (mutated in place), which surfaces as a ``when``-only row
    in the variants.csv: the customer fills the condition to show/hide the
    section, or leaves it blank for a plain unconditional loop. A loop or
    region name colliding with a ``[[$token]]`` key raises :class:`ValueError`.

    **Conditional regions** (``[Name?]`` opener) wrap the enclosed content —
    typically one or more table rows — in a presence/condition guard instead of
    a ``#foreach``. Three cases:

    - **inside a loop** → ``#if($<iterator>.<Name>)`` … ``#end`` emitted
      directly (per-item coverage presence; survives Legs 2/3 untouched, no
      customer fill, no plugin key). Warned when ``Name`` is not a registry
      coverage of that exposure.
    - **document level, Name is a registry coverage** → ``#if($doc.<Name>)``
      scaffold + a ``render: template`` block carrying ``presence`` metadata;
      the variants.csv skips it and Leg 4 auto-emits the presence Boolean
      (any item carries the coverage). No customer fill.
    - **document level, any other Name** → ``#if($doc.<Name>)`` scaffold + a
      plain ``render: template`` block: a ``when``-only variants.csv row the
      customer fills (blank = always render, guard stripped by Leg 3).

    A ``[[$token]]`` placeholder inside a loop is allowed but warned —
    conditions are document-scoped, so it renders identically for every item.
    Legacy ``[Name]`` openers (no slash) draw a migration warning when their
    ``[/Name]`` closer exists; unmatched markers warn and stay literal text.
    """
    pairs: list[tuple[int, int, int, int, str]] = []
    region_pairs: list[tuple[int, int, int, int, str]] = []
    pending: dict[str, tuple[int, int, str]] = {}
    legacy_openers: set[str] = set()
    for m in _LOOP_MARKER_RE.finditer(html):
        leading, name, trailing = m.group(1) == "/", m.group(2), m.group(3)
        if leading and trailing:
            print(
                f"WARNING: malformed marker [/{name}{trailing}] — left as literal text",
                file=sys.stderr,
            )
        elif not leading and trailing:  # [Name/] loop or [Name?] region opener
            kind = "loop" if trailing == "/" else "region"
            if name in pending:
                print(
                    f"WARNING: [{name}{trailing}] reopened before [/{name}] — earlier opener ignored",
                    file=sys.stderr,
                )
            pending[name] = (m.start(), m.end(), kind)
        elif leading:  # [/Name] closer
            if name in pending:
                o_start, o_end, kind = pending.pop(name)
                dest = pairs if kind == "loop" else region_pairs
                dest.append((o_start, o_end, m.start(), m.end(), name))
            elif name in legacy_openers:
                legacy_openers.discard(name)
                print(
                    f"WARNING: legacy loop syntax [{name}]…[/{name}] — the opener is now "
                    f"[{name}/]; markers left as literal text",
                    file=sys.stderr,
                )
            else:
                print(f"WARNING: [/{name}] closer without opener — left as literal text", file=sys.stderr)
        else:  # plain [Name] — legacy opener candidate; silent unless a closer follows
            legacy_openers.add(name)
    for name, (_s, _e, kind) in pending.items():
        marker = f"[{name}/]" if kind == "loop" else f"[{name}?]"
        print(f"WARNING: {marker} never closed — left as literal text", file=sys.stderr)
    if not pairs and not region_pairs:
        return html, fields, []

    # A [[$token]] inside a loop renders identically for every item (conditions
    # are document-scoped) — allowed, but worth a warning.
    for o_start, o_end, c_start, _c_end, name in pairs:
        if re.search(r"\$doc\.\w+", html[o_end:c_start]):
            print(
                f"WARNING: [[$token]] block(s) inside loop [{name}/] are document-scoped — "
                "the same text renders for every item",
                file=sys.stderr,
            )

    # Classify each [Name?] region: inside a loop → direct iterator guard;
    # document level → #if($doc.<Name>) + a render: template block, carrying
    # presence metadata when the name is a registry coverage (auto-wired by
    # Leg 4, skipped in the variants.csv).
    cov_map, loop_iters = _registry_coverage_map(registry_path) if region_pairs else ({}, {})
    region_guards: dict[tuple[int, int], str] = {}  # (o_start, c_start) -> guard ref
    doc_level_regions: list[tuple[str, dict | None]] = []  # (name, coverage info|None)
    loop_value_regions: list[tuple[str, str, str]] = []  # (name, loop_name, iterator)
    for ro_start, _ro_end, rc_start, _rc_end, name in region_pairs:
        enclosing = None  # innermost loop containing this region
        for o_start, o_end, c_start, _c_end, lname in pairs:
            if o_end <= ro_start and rc_start <= c_start:
                if enclosing is None or o_start > enclosing[0]:
                    enclosing = (o_start, lname)
        if enclosing is not None:
            loop_name = enclosing[1]
            info = cov_map.get(name) or {}
            is_loop_coverage = info.get("exposure") == loop_name
            # Use the registry iterable's iterator so the guard references the
            # same variable Leg 3's #foreach directive declares; name.lower()
            # only as the registry-less fallback.
            it_var = (
                info["iterator"] if is_loop_coverage and info.get("iterator")
                else loop_iters.get(loop_name) or loop_name.lower()
            )
            if cov_map and not is_loop_coverage:
                # In-loop VALUE region: not a coverage of this exposure, so the
                # rows show/hide on a per-item condition the customer fills as a
                # when-only variants.csv row (paths rooted at the iterator).
                # Leg 3 compiles the condition to an in-template #if inside the
                # loop; the plugin never sees it (per-item, not doc-scoped).
                print(
                    f"NOTE: [{name}?] inside loop [{loop_name}/] — `{name}` is not a "
                    f"registry coverage of `{loop_name}`; fill its variants.csv `when` "
                    f"with a per-item condition rooted at `{it_var}` (blank = always render)",
                    file=sys.stderr,
                )
                region_guards[(ro_start, rc_start)] = f"$doc.{name}"
                loop_value_regions.append((name, loop_name, it_var))
            else:
                # Coverage presence (or registry-less fallback): guard directly.
                region_guards[(ro_start, rc_start)] = f"${it_var}.{name}"
        else:
            region_guards[(ro_start, rc_start)] = f"$doc.{name}"
            doc_level_regions.append((name, cov_map.get(name)))

    # Every loop gets a when-only variants.csv row: append a render: template
    # block keyed by the loop name (blank `when` = unconditional loop).
    # Document-level [Name?] regions get one too — with `presence` metadata
    # (no customer fill) when the name is a registry coverage.
    if cond_blocks is not None:
        taken = {b["key"] for b in cond_blocks}
        next_id = max((b.get("id") or 0 for b in cond_blocks), default=0) + 1
        block_specs = [(p[4], None, "loop") for p in pairs] + [
            (name, cov_info, "region") for name, cov_info in doc_level_regions
        ] + [
            (name, {"loop": loop_name, "iterator": it_var}, "loop_region")
            for name, loop_name, it_var in loop_value_regions
        ]
        for name, meta, kind in block_specs:
            if name in taken:
                marker = f"[{name}/]" if kind == "loop" else f"[{name}?]"
                raise ValueError(
                    f"{marker} collides with another block keyed `{name}` — "
                    "loop, region and conditional-token names share one key space; rename one"
                )
            taken.add(name)
            block = {
                "id": next_id,
                "key": name,
                "placeholder": None,
                "variant": False,
                "render": "template",
                "source_text": "",
                "top_level": True,
            }
            if kind == "region" and meta:
                block["presence"] = {"coverage": name, **meta}
            elif kind == "loop_region":
                block["loop_scope"] = meta
            cond_blocks.append(block)
            next_id += 1

    loop_field_lists: dict[str, list[dict]] = {p[4]: [] for p in pairs}
    top_fields: list[dict] = []
    for f in fields:
        # Reject a following word char (partial-name match) or a `.word`
        # path continuation (a shorter token sitting inside a longer dotted
        # accessor), but ALLOW a trailing sentence period: `...partsCovered.`
        # must still match, else the field falls through to document scope and
        # loses its loop/coverage guard.
        token_re = re.compile(re.escape(f["token"]) + r"(?!\w|\.\w)")
        positions = [t.start() for t in token_re.finditer(html)]
        target = None
        for o_start, o_end, c_start, _c_end, name in pairs:
            if positions and all(o_end <= p < c_start for p in positions):
                target = name
                break
        if target is None:
            top_fields.append(f)
        else:
            loop_field_lists[target].append(f)

    # Replace marker spans in reverse offset order so earlier offsets stay valid.
    # Each loop is wrapped in its own #if($doc.<Name>) guard — Leg 3 resolves it
    # to #if($data.<Name>) when the customer filled the loop's `when` row, or
    # strips it when the row was left blank (unconditional loop). A [Name?]
    # region opener becomes a single #if — iterator-scoped ($item.<Cov>, final
    # as emitted) or document-scoped ($doc.<Name>, resolved/stripped by Leg 3).
    spans: list[tuple[int, int, str]] = []
    for o_start, o_end, c_start, c_end, name in pairs:
        spans.append((o_start, o_end, f"#if($doc.{name})\n#foreach (${name.lower()} in $TBD_{name})"))
        spans.append((c_start, c_end, "#end\n#end"))
    for ro_start, ro_end, rc_start, rc_end, _name in region_pairs:
        spans.append((ro_start, ro_end, f"#if({region_guards[(ro_start, rc_start)]})"))
        spans.append((rc_start, rc_end, "#end"))
    for start, end, text in sorted(spans, reverse=True):
        html = html[:start] + text + html[end:]
    for pattern, repl in _DIRECTIVE_LINE_CLEANUP:
        html = pattern.sub(repl, html)

    loops = [
        {
            "name": name,
            "token": f"$TBD_{name}",
            "iterator": f"${name.lower()}",
            "fields": loop_field_lists[name],
        }
        for _os, _oe, _cs, _ce, name in pairs
    ]
    return html, top_fields, loops


def _normalise_for_leg2(fields: list[dict], source_name: str, loops: list[dict] | None = None) -> dict:
    """Convert fields list to a leg2-compatible .mapping.yaml dict (uses placeholder)."""
    import datetime as dt

    def _variable(f: dict, loop_name: str | None = None) -> dict:
        context: dict = {
            "parent_tag": "p",
            "line": None,
            "nearest_label": "",
        }
        if loop_name:
            context["loop"] = loop_name
        var = {
            "name": f["name"],
            "placeholder": f["token"],
            "type": "loop_field" if loop_name else "variable",
            "context": context,
            "data_source": f.get("data_source", ""),
            "occurrence": f.get("occurrence", "required"),
        }
        if f.get("candidate"):
            var["candidate"] = f["candidate"]
        return var

    loop_entries = [
        {
            "name": loop["name"],
            "placeholder": loop["token"],
            "type": "loop",
            "iterator": loop["iterator"],
            "detection": "marker",
            "context": {},
            "data_source": "",
            "fields": [_variable(f, loop_name=loop["name"]) for f in loop["fields"]],
        }
        for loop in (loops or [])
    ]
    return {
        "schema_version": "1.0",
        "source": source_name,
        "generated_at": dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%S+00:00"),
        "variables": [_variable(f) for f in fields],
        "loops": loop_entries,
    }


def write_leg2_mapping(
    fields: list[dict], source_name: str, output_path: Path, loops: list[dict] | None = None
) -> None:
    """Write {stem}.mapping.yaml — leg2-compatible format with placeholder field."""
    data = _normalise_for_leg2(fields, source_name, loops=loops)
    validate_contract(data, MappingDoc, artifact="mapping.yaml", path=output_path)
    output_path.write_text(
        yaml.dump(data, default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )


# ---------------------------------------------------------------------------
# Write conditional form (E-T4)
# ---------------------------------------------------------------------------

def _braces_to_tbd(text: str) -> str:
    """Convert {name} → $TBD_name (canonical machine form).

    No-op on text already in $TBD_ form, so forms written before the {field}
    display change still parse. Occurrence symbols ({$x}, {+x}, {*x}) are
    accepted and dropped — the canonical token never carries them.
    """
    return _FIELD_RE.sub(lambda m: "$TBD_" + m.group(2), text)


def write_variants_csv(blocks: list[dict], stem: str, output_path: Path) -> None:
    """Write ``{stem}.variants.csv`` — the single human-fill file for *all*
    conditional blocks (variants-only plan §2.1).

    One row group per block, keyed by the block's join ``key`` (pre-filled — the
    customer never renames it). The stub lists the blocks with **no fabricated
    conditions** — every ``when`` is blank for the customer to fill. Block kind
    drives the row shape:

    - **variant** (``[[$token]]``): one conditioned row + one default row, both
      blank; the customer fills/adds rows and supplies each variant's ``text``.
    - **template** (``render: template`` — a ``[Name/]`` loop section): a single
      ``when``-only row, ``text`` left blank because the section's wording stays
      in the document. The customer fills the ``when`` to show/hide the section,
      or leaves it blank for a plain unconditional loop.

    Block metadata that the three columns can't carry (id, ``render`` flag,
    ``source_text``, nesting) travels in the machine sidecar
    ``{stem}.conditional-blocks.yaml`` (see :func:`write_conditional_blocks`),
    which the parse step reads back alongside this CSV.

    One genuinely-unsupported edge (documented, not handled): an N-way block
    whose variants each carry their *own* loop — loop bodies can't live in a CSV
    ``text`` cell and ``render: template`` is binary show/hide, not N-way.
    """
    import csv  # noqa: PLC0415
    import io  # noqa: PLC0415
    if not blocks:
        return
    rows: list[list[str]] = []
    for b in blocks:
        key = block_key(b)
        if b.get("presence"):
            # [Name?] coverage-presence region: auto-wired by Leg 4 (any item
            # carries the coverage) — nothing for the customer to fill.
            continue
        if b.get("render") == "template":
            # [Name/] loop section or [Name?] conditional region: a single
            # `when`-only row (blank = unconditional / always render).
            rows.append([key, "", ""])
        else:
            # N-way [[$token]]: one conditioned row to fill + the default row.
            rows.append([key, "", ""])
            rows.append([key, "", ""])
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["placeholder", "when", "text"])
    w.writerows(rows)
    content = buf.getvalue()

    # Clobber guard (human-fill friction Gap 2) + orphan purge: a full re-ingest
    # must not destroy a customer's filled conditions, but must also reconcile the
    # file to the CURRENT document — otherwise a stem reused across documents (or a
    # marker deleted from the doc) leaves orphan rows the author edits for markers
    # that no longer exist (the CGL "Pet" residue). Reconcile against this run's
    # block keys: keep hand-edited rows whose block still exists, DROP orphans whose
    # block is gone, APPEND stubs for genuinely new blocks.
    if output_path.exists():
        existing = output_path.read_text(encoding="utf-8")
        if existing != content:
            current_keys = {r[0] for r in rows}
            existing_rows = [r for r in csv.reader(io.StringIO(existing)) if r]
            if existing_rows and existing_rows[0][:1] == ["placeholder"]:
                existing_rows = existing_rows[1:]  # drop the header row
            kept = [r for r in existing_rows if r[0] in current_keys]
            orphans = sorted({r[0] for r in existing_rows if r[0] not in current_keys})
            existing_keys = {r[0] for r in existing_rows}
            new_rows = [r for r in rows if r[0] not in existing_keys]
            if not orphans and not new_rows:
                print(
                    f"WARN: {output_path.name} already exists with edits — keeping it, "
                    "NOT overwriting (delete it first to regenerate the blank stub).",
                    file=sys.stderr,
                )
                return
            buf = io.StringIO()
            w = csv.writer(buf)
            w.writerow(["placeholder", "when", "text"])
            w.writerows(kept + new_rows)
            output_path.write_text(buf.getvalue(), encoding="utf-8")
            msgs = []
            if new_rows:
                msgs.append(f"appended stub(s) for new block(s): {sorted({r[0] for r in new_rows})}")
            if orphans:
                msgs.append(f"dropped orphan row(s) for block(s) no longer in the document: {orphans}")
            print(f"NOTE: {output_path.name} has hand-edits — kept them; " + "; ".join(msgs),
                  file=sys.stderr)
            return
    output_path.write_text(content, encoding="utf-8")


# Fields persisted in the machine sidecar (everything the 3-column CSV can't carry).
def write_conditional_blocks(blocks: list[dict], output_path: Path) -> None:
    """Write ``{stem}.conditional-blocks.yaml`` — the machine sidecar that
    carries per-block metadata the human ``variants.csv`` cannot (variants-only
    plan §2.1, machine-sidecar design). The parse step reads it back alongside
    the filled CSV to rebuild ``conditional-registry.yaml`` now that
    ``conditional-form.md`` is retired.
    """
    payload = [
        {
            "id": b.get("id"),
            "key": block_key(b),
            "placeholder": b.get("placeholder"),
            "variant": bool(b.get("variant")),
            "render": b.get("render", "plugin"),
            "source_text": b.get("source_text", ""),
            "top_level": b.get("top_level", True),
            "parent_id": b.get("parent_id"),
            "depth": b.get("depth", 0),
            **({"presence": b["presence"]} if b.get("presence") else {}),
            **({"loop_scope": b["loop_scope"]} if b.get("loop_scope") else {}),
        }
        for b in blocks
    ]
    output_path.write_text(
        yaml.dump(payload, default_flow_style=False, allow_unicode=True), encoding="utf-8"
    )


def load_conditional_blocks(path: Path) -> list[dict]:
    """Read the machine sidecar written by :func:`write_conditional_blocks`."""
    data = yaml.safe_load(Path(path).read_text(encoding="utf-8")) or []
    return data if isinstance(data, list) else []


# ---------------------------------------------------------------------------
# Parse filled conditional form → conditional-registry.yaml (E-T5)
# ---------------------------------------------------------------------------

# A variant pointer block in the form carries "Variant placeholder: `$token`"
# instead of a Condition: line (its rows live in the sibling .variants.csv).
_VARIANT_FORM_RE = re.compile(
    r"##\s+Block\s+(\d+)\s*\n+>\s+(.+?)\s*\n+Variant placeholder:\s*`\$([A-Za-z_]\w*)`",
    re.DOTALL,
)


def parse_conditional_form(
    md_path: Path,
    *,
    variants_csv: Path | None = None,
    registry: dict | None = None,
    classpath: str | None = None,
    product: str | None = None,
) -> list[dict]:
    """Parse a customer-filled conditional form. Returns list of block dicts.

    Binary and template blocks parse from their ``Condition:`` line as before.
    Variant blocks (``Variant placeholder: `$token```) carry no condition here —
    their rows come from the sibling ``<stem>.variants.csv`` (or ``variants_csv``
    override), normalised via :func:`condition_dsl.parse_variants_csv` and merged
    in by placeholder. Any CSV/DSL validation error raises :class:`ValueError`
    so the caller never writes a half-valid registry (§1a / Phase 3).
    """
    from velocity_converter.condition_dsl import parse_variants_csv  # noqa: PLC0415

    text = md_path.read_text(encoding="utf-8")
    blocks = []

    # The body capture is *tempered* so it cannot run past the next "## Block"
    # header or a "Variant placeholder:" line — without this, a variant block
    # (which has no Condition: line) would let the binary match swallow the
    # following binary block and steal its condition (duplicate-id collision).
    block_re = re.compile(
        r"##\s+Block\s+(\d+)\s*\n+>\s+"
        r"((?:(?!\n##\s+Block\s+\d+|\nVariant placeholder:).)+?)\s*\n+"
        r"(Rendering:\s*template[^\n]*\n+)?Condition:\s*([^\n]*)",
        re.DOTALL,
    )
    for m in block_re.finditer(text):
        block_id = int(m.group(1))
        source_text = _braces_to_tbd(m.group(2).strip())
        raw_condition = m.group(4).strip()
        # Take only the first line of the condition (customer may add notes below)
        condition_line = raw_condition.splitlines()[0].strip() if raw_condition else ""
        conditions = [condition_line] if condition_line else []
        block = {
            "id": block_id,
            "source_text": source_text,
            "conditions": conditions,
            "operator": "AND",
        }
        if m.group(3):
            block["render"] = "template"
        blocks.append(block)

    # Variant pointer blocks (no Condition: line — filled via the CSV).
    variant_blocks: dict[str, dict] = {}
    for m in _VARIANT_FORM_RE.finditer(text):
        block_id = int(m.group(1))
        placeholder = m.group(3)
        block = {
            "id": block_id,
            "key": placeholder,
            "placeholder": placeholder,
            "variant": True,
            "source_text": _braces_to_tbd(m.group(2).strip()),
        }
        blocks.append(block)
        variant_blocks[placeholder] = block

    blocks.sort(key=lambda b: b["id"])

    # Merge the sibling variants.csv into the variant blocks.
    csv_path = variants_csv
    if csv_path is None:
        stem = md_path.name
        if stem.endswith(".conditional-form.md"):
            stem = stem[: -len(".conditional-form.md")]
        sibling = md_path.parent / f"{stem}.variants.csv"
        if sibling.is_file():
            csv_path = sibling

    if variant_blocks and csv_path is None:
        raise ValueError(
            f"the form declares variant block(s) {sorted(variant_blocks)} but no "
            f"variants CSV was found next to {md_path.name} — provide --variants-csv"
        )

    if csv_path is not None:
        result = parse_variants_csv(csv_path, registry, classpath=classpath, product=product)
        errors = list(result.errors)
        for ph, data in result.placeholders.items():
            b = variant_blocks.get(ph)
            if b is None:
                errors.append(
                    f"variants CSV placeholder '{ph}' has no matching [[${ph}]] block in the form"
                )
                continue
            b["variants"] = data["variants"]
            b["default"] = data["default"]
            b["scope"] = data["scope"]
        for ph, b in variant_blocks.items():
            if ph not in result.placeholders:
                errors.append(f"variant block [[${ph}]] has no rows in the variants CSV")
        if errors:
            raise ValueError(
                "variant CSV validation failed (registry NOT written):\n  - "
                + "\n  - ".join(errors)
            )

    return blocks


# A nested reference in a parsed variant text cell: condition_dsl peels [[$x]] → $doc.x
# before storing, so the join with sibling blocks works off the machine form.
_DOC_REF_RE = re.compile(r"\$doc\.([A-Za-z_]\w*)")


def _nested_refs(data: dict) -> set[str]:
    """Placeholder keys referenced via ``$doc.<key>`` in a placeholder's texts."""
    texts = [data.get("default") or ""]
    texts += [v.get("text") or "" for v in (data.get("variants") or [])]
    return set(_DOC_REF_RE.findall(" ".join(texts)))


def _find_ref_cycle(refs: dict[str, set[str]]) -> list[str] | None:
    """Return a cycle path through the nested-ref graph, or None. DFS 3-colour."""
    WHITE, GREY = 0, 1
    color: dict[str, int] = {}
    stack: list[str] = []

    def dfs(u: str) -> list[str] | None:
        color[u] = GREY
        stack.append(u)
        for v in refs.get(u, ()):
            if v not in refs:  # missing referent — reported separately
                continue
            if color.get(v) == GREY:
                return stack[stack.index(v):] + [v]
            if color.get(v, WHITE) == WHITE:
                cyc = dfs(v)
                if cyc:
                    return cyc
        stack.pop()
        color[u] = 2
        return None

    for ph in refs:
        if color.get(ph, WHITE) == WHITE:
            cyc = dfs(ph)
            if cyc:
                return cyc
    return None


def parse_variants_csv_to_blocks(
    csv_path: Path,
    blocks_meta: list[dict],
    registry: dict | None = None,
    *,
    classpath: str | None = None,
    product: str | None = None,
    doc_scope: str | None = None,
) -> list[dict]:
    """Merge a filled ``<stem>.variants.csv`` with its machine sidecar
    (``blocks_meta`` from :func:`load_conditional_blocks`) into the block list
    written to ``conditional-registry.yaml`` (variants-only plan §2.3).

    This is the single parse path — every block kind flows through one CSV:
      - **variant** (``[[$token]]``) → ``variants`` + ``default`` + ``scope``.
      - **template** (``render: template`` — a ``[Name/]`` loop row) → a Boolean:
        the single ``when`` is carried as a one-entry ``variants`` payload (text
        blank) and evaluated by Leg 4 through the same DSL→Java path as N-way
        blocks. A blank ``when`` means the loop is unconditional — the block is
        **omitted** from the registry entirely (Leg 3 strips its ``#if`` guard,
        Leg 4 puts nothing).

    Block metadata (id, ``render``, ``source_text``, nesting) comes from the
    sidecar; the human CSV only supplies ``when``/``text``. Any CSV/DSL/scope
    validation error raises :class:`ValueError` so a half-valid registry is never
    written.
    """
    from velocity_converter.condition_dsl import parse_variants_csv  # noqa: PLC0415

    template_phs = {block_key(b) for b in blocks_meta if b.get("render") == "template"}
    by_key = {block_key(b): b for b in blocks_meta}
    loop_scoped = {
        block_key(b): str((b.get("loop_scope") or {}).get("iterator") or "")
        for b in blocks_meta
        if b.get("loop_scope")
    }
    result = parse_variants_csv(
        csv_path, registry, classpath=classpath, product=product,
        template_placeholders=template_phs, doc_scope=doc_scope,
        loop_scoped=loop_scoped or None,
    )
    errors = list(result.errors)

    # Nested [[$x]] refs (peeled to $doc.x at parse): validate the reference graph
    # before building blocks so a half-valid registry is never written.
    phs = result.placeholders
    refs = {ph: _nested_refs(data) for ph, data in phs.items()}
    for ph, deps in refs.items():
        ph_scope = phs[ph].get("scope") or ""
        for d in deps:
            if d == ph:
                errors.append(f"'{ph}' references itself via [[${ph}]] — nesting cannot be self-referential")
                continue
            if d not in phs:
                errors.append(f"'{ph}' references [[${d}]] but the variants CSV has no row '{d}'")
                continue
            # A conditional label must share its parent's scope; an unconditional
            # one (scope "") composes into either overload, so it is exempt.
            d_scope = phs[d].get("scope") or ""
            if d_scope and ph_scope and d_scope != ph_scope:
                errors.append(
                    f"'{ph}' ({ph_scope}-scoped) references [[${d}]] which is {d_scope}-scoped — "
                    "a nested label must share its parent's scope (or be unconditional)"
                )
    cycle = _find_ref_cycle(refs)
    if cycle:
        errors.append("nested [[$...]] reference cycle: " + " → ".join(cycle))

    out_blocks: list[dict] = []
    for b in blocks_meta:
        key = block_key(b)
        block: dict = {
            "id": b.get("id"),
            "key": key,
            "source_text": b.get("source_text", ""),
            "parent_id": b.get("parent_id"),
            "depth": b.get("depth", 0),
        }
        if b.get("presence"):
            # [Name?] coverage-presence region: no CSV rows (nothing to fill) —
            # register directly; Leg 4 emits the presence Boolean, Leg 3 renames
            # the #if($doc.<key>) guard to #if($data.<key>).
            block["render"] = "template"
            block["presence"] = b["presence"]
            block["variants"] = []
            block["default"] = None
            out_blocks.append(block)
            continue
        data = result.placeholders.get(key)
        if data is None:
            errors.append(f"block '{key}' has no rows in the variants CSV")
            out_blocks.append(block)
            continue
        if b.get("render") == "template":
            if not data["variants"]:
                # Blank `when` = unconditional loop: no registry entry at all —
                # Leg 3 strips the #if($doc.<key>) guard, Leg 4 puts nothing.
                continue
            block["render"] = "template"
            block["scope"] = data["scope"]
            # The single `when` rides in a one-entry variants payload (text blank);
            # Leg 4's template branch reads variants[0].when and emits a Boolean.
            block["variants"] = data["variants"]
            block["default"] = None
            if b.get("loop_scope"):
                # In-loop value region: Leg 3 compiles the `when` to an
                # in-template #if inside the loop; Leg 4 skips it (per-item).
                block["loop_scope"] = b["loop_scope"]
        else:
            block["placeholder"] = b.get("placeholder")
            block["scope"] = data["scope"]
            block["variants"] = data["variants"]
            block["default"] = data["default"]
        out_blocks.append(block)

    # Placeholders with no sidecar block: synthesize one if it is a nested-only label
    # (referenced via [[$x]] from another row), else it is a true orphan — keep the
    # error. Synthesized blocks carry no document marker (source_text "") and compose
    # inside their referrer's plugin string; the template never sees them.
    all_refs: set[str] = set().union(*refs.values()) if refs else set()
    next_id = max((b.get("id") or 0 for b in blocks_meta), default=0) + 1
    for ph in sorted(phs):
        if ph in by_key:
            continue
        if ph in all_refs:
            data = phs[ph]
            out_blocks.append({
                "id": next_id,
                "key": ph,
                "source_text": "",
                "parent_id": None,
                "depth": 0,
                "placeholder": ph,
                "scope": data["scope"],
                "variants": data["variants"],
                "default": data["default"],
            })
            next_id += 1
        else:
            errors.append(
                f"variants CSV placeholder '{ph}' has no matching block in the sidecar "
                "(was it renamed?)"
            )

    if errors:
        raise ValueError(
            "variant CSV validation failed (registry NOT written):\n  - "
            + "\n  - ".join(errors)
        )
    return out_blocks


def _primary_root_scope(mapping_path: Path) -> str | None:
    """Condition scope ('quote'|'policy') of a mapping's primary rendering root.

    A quote-rooted document conditions on quote accessors; everything else
    (segment/policy/term) conditions on policy accessors. Returns None when the
    mapping is absent or carries no rendering root (caller stays scope-blind).
    """
    if not mapping_path.is_file():
        return None
    try:
        data = yaml.safe_load(mapping_path.read_text(encoding="utf-8")) or {}
    except Exception:
        return None
    roots = data.get("rendering_roots") or []
    if not roots:
        return None
    primary = next((r for r in roots if r.get("primary")), roots[0])
    return "quote" if str(primary.get("id") or "").lower() == "quote" else "policy"


def _product_from_mapping(mapping_path: Path) -> str | None:
    """Product name from a mapping.yaml ('product' key), or None. Used to form the
    FQCN for optional jar verification of condition paths."""
    if not mapping_path.is_file():
        return None
    try:
        data = yaml.safe_load(mapping_path.read_text(encoding="utf-8")) or {}
    except Exception:
        return None
    return (data.get("product") or "").strip() or None


def write_conditional_registry(blocks: list[dict], output_path: Path) -> None:
    """Write {stem}.conditional-registry.yaml for Leg 4."""
    validate_contract(
        blocks, ConditionalRegistry, artifact="conditional-registry.yaml", path=output_path
    )
    output_path.write_text(
        yaml.dump(blocks, default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )


# ---------------------------------------------------------------------------
# Output writer
# ---------------------------------------------------------------------------

def write_output(html: str, input_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{input_path.stem}.raw.html"
    output_path.write_text(html, encoding="utf-8")
    print(f"Wrote {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Document parse (shared by full ingest and --scan)
# ---------------------------------------------------------------------------

@dataclass
class ParseResult:
    """Everything a document parse yields, before any artifact is written.

    Both the full ingest and the lightweight ``--scan`` mode call
    :func:`_parse_document` and then choose which of these to persist: full
    ingest writes the machine artifacts (raw/annotated HTML + mapping) plus the
    human-fill files; ``--scan`` writes *only* the human-fill files
    (conditional-form + variants.csv) so the customer can start filling them
    while the rest of the pipeline is deferred.
    """

    stem: str
    raw_html: str
    annotated: str
    fields: list[dict] = field(default_factory=list)
    loops: list[dict] = field(default_factory=list)
    blocks: list[dict] = field(default_factory=list)


def _parse_document(
    input_path: Path,
    *,
    path_map: Path | None = None,
    registry_path: Path | None = None,
    converter: str = "soffice",
) -> ParseResult:
    """Convert + extract a .docx/.pdf into a :class:`ParseResult`, writing nothing.

    This is the full Leg 0 markup parse — convert → fields → conditionals →
    loops — factored out so the full ingest and ``--scan`` share one code path
    (a re-parse across two separate invocations is deterministic and cheap, so
    nothing is cached). The source document is never modified; an optional
    Leg -1 ``path_map`` is applied to the working HTML before extraction.
    """
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        raw_html = (
            convert_docx(input_path) if converter == "legacy"
            else convert_docx_soffice(input_path)
        )
    else:
        raw_html = convert_pdf(input_path)

    # Apply a Leg -1 path-map (bare {leaf} → full accessor) if supplied, so the
    # author's friendly leaves resolve before extraction.
    if path_map is not None:
        raw_html = apply_path_map(raw_html, path_map)

    # Rendering root from the filename suffix (e.g. ...(segment)) — drives the
    # $data.<root> entity-key splice so resolved scalars match renderingData.
    from velocity_converter.leg2_fill_mapping import parse_rendering_roots  # noqa: PLC0415
    _roots, _ = parse_rendering_roots(input_path.stem)
    rendering_root = _roots[0] if _roots else None

    # Extract + annotate fields, then conditionals.
    fields = extract_fields(raw_html, registry_path=registry_path, rendering_root=rendering_root)
    annotated = annotate_fields(raw_html, fields)
    blocks = extract_conditionals(annotated)
    annotated = annotate_conditionals(annotated, blocks)

    # Extract [Name]...[/Name] loop sections (after field annotation so loop
    # membership is decided on $TBD_* token positions). Blocks containing a
    # loop flip to render: template — this is the only place the conditional
    # form's `render:` note is set, so --scan must run it too.
    annotated, fields, loops = extract_loops(
        annotated, fields, cond_blocks=blocks, registry_path=registry_path
    )

    return ParseResult(
        stem=input_path.stem,
        raw_html=raw_html,
        annotated=annotated,
        fields=fields,
        loops=loops,
        blocks=blocks,
    )


def _write_human_fill_files(blocks: list[dict], stem: str, output_dir: Path) -> list[Path]:
    """Write the single customer-facing hand-fill file for a parse's conditional
    blocks (variants-only plan §2.1): ``{stem}.variants.csv`` in the flat
    ``action-needed/`` space, covering *every* block kind (binary/template/
    variant). The machine sidecar that pairs with it is written separately by
    the full ingest (see :func:`write_conditional_blocks`). Returns the paths
    written (empty when there are no conditional blocks).
    """
    if not blocks:
        print("No [[conditional]] blocks found — skipping variants CSV.")
        return []
    csv_path = action_needed_file(output_dir, f"{stem}.variants.csv")
    write_variants_csv(blocks, stem, csv_path)
    print(f"Wrote {csv_path}")
    return [csv_path]


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _discover_registry(start: Path) -> Path | None:
    """Walk up from ``start`` looking for a path-registry.yaml (shared by the
    convert and parse-conditional-form modes)."""
    cur = start.resolve()
    for _ in range(8):
        for rel in ("registry/path-registry.yaml", "path-registry.yaml"):
            candidate = cur / rel
            if candidate.is_file():
                return candidate
        if cur.parent == cur:
            break
        cur = cur.parent
    return None


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Leg 0: convert PDF/Word to raw HTML and extract fields/conditionals."
    )
    parser.add_argument("--input", default=None, help="Path to .docx or .pdf file")
    parser.add_argument("--output-dir", default=None, help="Output directory (default: input's parent)")
    parser.add_argument(
        "--scan",
        action="store_true",
        help="Scan mode: parse the document and emit ONLY the human-fill file "
        "(variants.csv), no machine artifacts — front-loads the customer handoff "
        "before the full ingest runs.",
    )
    parser.add_argument(
        "--path-map",
        default=None,
        metavar="PATH_MAP.yaml",
        help="Apply a Leg -1 path-map (bare {leaf} → full accessor) before extraction",
    )
    parser.add_argument(
        "--converter",
        choices=["soffice", "legacy"],
        default="soffice",
        help=".docx→HTML converter: 'soffice' (default — LibreOffice headless, "
        "preserves the document's styling; requires LibreOffice installed) or "
        "'legacy' (structure-only python-docx converter, styling discarded).",
    )
    parser.add_argument(
        "--parse-variants-csv",
        default=None,
        metavar="FILLED.variants.csv",
        help="Parse a filled variants.csv (+ its machine sidecar "
        "<stem>.conditional-blocks.yaml) → conditional-registry.yaml. The single "
        "parse path in the variants-only flow.",
    )
    parser.add_argument(
        "--parse-conditional-form",
        default=None,
        metavar="FILLED_FORM.md",
        help="LEGACY: parse an old filled conditional-form.md (+ sibling "
        "variants.csv) → conditional-registry.yaml. Retained for in-flight forms; "
        "new documents use --parse-variants-csv.",
    )
    parser.add_argument(
        "--blocks-sidecar",
        default=None,
        metavar="BLOCKS.yaml",
        help="Override the conditional-blocks.yaml sidecar location "
        "(default: <machine-dir>/<stem>.conditional-blocks.yaml)",
    )
    parser.add_argument(
        "--variants-csv",
        default=None,
        metavar="VARIANTS.csv",
        help="Override the variants CSV location (legacy form parse: default is the "
        "sibling <stem>.variants.csv)",
    )
    parser.add_argument(
        "--no-registry",
        action="store_true",
        help="Config-agnostic mode: skip registry auto-discovery everywhere "
        "(ingest classification and variants-CSV parse). Condition validation "
        "degrades to syntax-only; loop/region classification uses the "
        "registry-less fallbacks.",
    )
    parser.add_argument(
        "--customer-jar",
        default=None,
        metavar="customer-config.jar",
        help="Customer config jar. With --datamodel-jar (+ --product), a condition "
        "path absent from the curated registry is accepted when it resolves against "
        "the real model — the jar is authority over the registry.",
    )
    parser.add_argument(
        "--datamodel-jar",
        default=None,
        metavar="core-datamodel.jar",
        help="Core datamodel jar, paired with --customer-jar for jar-verified paths.",
    )
    parser.add_argument(
        "--product",
        default=None,
        help="Product name for jar verification (default: read from <stem>.mapping.yaml).",
    )
    args = parser.parse_args()

    # --- Mode: parse filled variants.csv (+ sidecar) → conditional-registry ---
    if args.parse_variants_csv:
        csv_path = Path(args.parse_variants_csv)
        if not csv_path.exists():
            print(f"Error: file not found: {csv_path}", file=sys.stderr)
            return 1
        name = csv_path.name
        stem = name[: -len(".variants.csv")] if name.endswith(".variants.csv") else csv_path.stem
        # Machine artifacts (sidecar + registry) live in the per-stem machine dir,
        # not next to the human CSV in action-needed/.
        machine_dir = (
            Path(args.output_dir) if args.output_dir
            else machine_dir_for_action_file(csv_path) or csv_path.parent
        )
        machine_dir.mkdir(parents=True, exist_ok=True)
        blocks_path = (
            Path(args.blocks_sidecar) if args.blocks_sidecar
            else machine_dir / f"{stem}.conditional-blocks.yaml"
        )
        if not blocks_path.exists():
            print(
                f"Error: machine sidecar not found: {blocks_path}\n"
                "  Run the full Leg 0 ingest first (it writes the sidecar that pairs "
                "with the variants CSV).",
                file=sys.stderr,
            )
            return 1

        from velocity_converter.condition_dsl import load_registry_dict  # noqa: PLC0415

        reg_path = None if args.no_registry else (
            _discover_registry(csv_path.parent) or _discover_registry(machine_dir)
        )
        registry = load_registry_dict(reg_path) if reg_path else None
        # Resolve bare-leaf conditions against the document's rendering root, so a
        # (quote) document conditions on quote.data.<f>, not the policy.data home.
        doc_scope = _primary_root_scope(machine_dir / f"{stem}.mapping.yaml")
        # Optional jar-as-authority: when a jar is supplied, a fully-qualified
        # condition path missing from the curated registry is accepted if it
        # resolves against the real model (verified via javap downstream).
        classpath = product = None
        if args.customer_jar and args.datamodel_jar:
            classpath = f"{Path(args.customer_jar).resolve()}:{Path(args.datamodel_jar).resolve()}"
            product = args.product or _product_from_mapping(machine_dir / f"{stem}.mapping.yaml")
        try:
            blocks_meta = load_conditional_blocks(blocks_path)
            blocks = parse_variants_csv_to_blocks(
                csv_path, blocks_meta, registry=registry, doc_scope=doc_scope,
                classpath=classpath, product=product,
            )
        except ValueError as exc:
            print(f"Error: {exc}", file=sys.stderr)
            return 1
        registry_path = machine_dir / f"{stem}.conditional-registry.yaml"
        write_conditional_registry(blocks, registry_path)
        print(f"Wrote {registry_path}")
        return 0

    # --- Mode: parse filled conditional form ---
    if args.parse_conditional_form:
        form_path = Path(args.parse_conditional_form)
        if not form_path.exists():
            print(f"Error: file not found: {form_path}", file=sys.stderr)
            return 1
        output_dir = Path(args.output_dir) if args.output_dir else form_path.parent
        output_dir.mkdir(parents=True, exist_ok=True)

        stem = form_path.stem
        if stem.endswith(".conditional-form"):
            stem = stem[: -len(".conditional-form")]

        from velocity_converter.condition_dsl import load_registry_dict  # noqa: PLC0415

        reg_path = None if args.no_registry else _discover_registry(form_path.parent)
        registry = load_registry_dict(reg_path) if reg_path else None
        try:
            blocks = parse_conditional_form(
                form_path,
                variants_csv=Path(args.variants_csv) if args.variants_csv else None,
                registry=registry,
            )
        except ValueError as exc:
            print(f"Error: {exc}", file=sys.stderr)
            return 1
        registry_path = output_dir / f"{stem}.conditional-registry.yaml"
        write_conditional_registry(blocks, registry_path)
        print(f"Wrote {registry_path}")
        return 0

    # --- Mode: convert document ---
    if not args.input:
        print("Error: --input is required (unless using --parse-variants-csv or "
              "--parse-conditional-form)", file=sys.stderr)
        return 1

    input_path = Path(args.input)

    if not input_path.exists():
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        return 1

    suffix = input_path.suffix.lower()

    if suffix == ".doc":
        print(
            "Error: legacy .doc format is not supported.\n"
            "Open the file in Microsoft Word → File → Save As → .docx, then re-run.",
            file=sys.stderr,
        )
        return 1

    if suffix not in (".docx", ".pdf"):
        print(
            f"Error: unsupported format '{suffix}'. Accepted: .docx, .pdf",
            file=sys.stderr,
        )
        return 1

    output_dir = Path(args.output_dir) if args.output_dir else input_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = input_path.stem

    # Find registry for explicit path resolution (walk up from input dir)
    registry_path = None if args.no_registry else _discover_registry(input_path.parent)

    # Validate the optional Leg -1 path-map before parsing. Source doc is untouched.
    path_map: Path | None = None
    if args.path_map:
        path_map = Path(args.path_map)
        if not path_map.exists():
            print(f"Error: path-map not found: {path_map}", file=sys.stderr)
            return 1

    # Single document parse — shared by --scan and full ingest. A ValueError is
    # an authoring error in the document (bare [[text]] block, duplicate token,
    # loop/token key collision) — report it cleanly, no traceback.
    try:
        pr = _parse_document(
            input_path, path_map=path_map, registry_path=registry_path,
            converter=args.converter,
        )
    except ValueError as exc:
        print(f"Error in {input_path.name}:\n{exc}", file=sys.stderr)
        return 1
    except RuntimeError as exc:
        # Token-integrity / conversion failure from the soffice path.
        print(f"Error in {input_path.name}:\n{exc}", file=sys.stderr)
        return 1

    # --- Sub-mode: scan — emit ONLY the human-fill files, defer machine artifacts.
    if args.scan:
        _write_human_fill_files(pr.blocks, stem, output_dir)
        return 0

    # Write raw HTML
    raw_path = output_dir / f"{stem}.raw.html"
    raw_path.write_text(pr.raw_html, encoding="utf-8")
    print(f"Wrote {raw_path}")

    # Write annotated HTML (pipeline input for Leg 1 / Leg 3)
    annotated_path = output_dir / f"{stem}.annotated.html"
    annotated_path.write_text(pr.annotated, encoding="utf-8")
    print(f"Wrote {annotated_path}")

    # Write leg2-compatible mapping
    mapping_path = output_dir / f"{stem}.mapping.yaml"
    write_leg2_mapping(pr.fields, f"{stem}.annotated.html", mapping_path, loops=pr.loops)
    print(f"Wrote {mapping_path}")

    # Write the machine sidecar (block metadata) that pairs with the human CSV.
    if pr.blocks:
        blocks_path = output_dir / f"{stem}.conditional-blocks.yaml"
        write_conditional_blocks(pr.blocks, blocks_path)
        print(f"Wrote {blocks_path}")

    # Write the single human-fill file (variants.csv — all conditional blocks).
    _write_human_fill_files(pr.blocks, stem, output_dir)

    return 0


if __name__ == "__main__":
    sys.exit(main())
