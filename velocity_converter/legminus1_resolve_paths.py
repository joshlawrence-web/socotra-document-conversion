#!/usr/bin/env python3
"""Leg -1: resolve bare ``{leaf}`` placeholders to full accessor paths.

Business problem: making document authors learn the exact accessor path
(``account.data.firstName``) before Leg 0 is too much. Leg -1 lets them write a
bare leaf (``{firstName}``) and does the registry lookup for them, emitting a
human-validated artifact *before* Leg 0 runs.

Modes::

    # 1. Suggest — doc → editable review + machine map + before/after audit
    python3 -m velocity_converter.legminus1_resolve_paths \
        --input <doc.docx|.pdf|.html> --registry registry/path-registry.yaml \
        --output-dir workspace/output

    # 2. Apply — fold the customer-filled CSV onto the canonical review → map + doc
    python3 -m velocity_converter.legminus1_resolve_paths \
        --parse-path-review-csv workspace/action-needed/<stem>.path-review.csv

    # 2b. Apply (operator) — parse the canonical .md directly (CSV not required)
    python3 -m velocity_converter.legminus1_resolve_paths \
        --parse-path-review workspace/output/<stem>/<stem>.path-review.md \
        --output-dir workspace/output/<stem>

Outputs:
    action-needed/<stem>.path-review.csv  — customer-fill: field / suggested / final
    <output-dir>/<stem>/<stem>.path-review.md   — canonical record (system copy)
    <output-dir>/<stem>/<stem>.path-map.yaml    — machine map (leaf → accessor) for Leg 0
    <output-dir>/<stem>/<stem>.path-changes.md  — before/after audit, one row per field
    <output-dir>/<stem>/<stem>.resolved.<ext>   — (apply) doc copy with full accessors

Matching is **registry-only** — no compiled SDK is consulted, so a "resolved"
leaf is registry-matched, not JAR-verified. Leg 2 still verifies paths against
the rendering root downstream.
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import re
import sys
from pathlib import Path

import yaml

from velocity_converter.leg0_ingest import (
    _FIELD_RE,
    _LOOP_MARKER_RE,
    convert_docx,
    convert_pdf,
)
from velocity_converter.models import OCCURRENCE_SYMBOLS
from velocity_converter.workspace import action_needed_file, machine_dir_for_action_file
from velocity_converter.registry_match import (
    build_candidate_index,
    match_leaf,
    parse_rendering_roots,
    suggest_loop_root,
)

_INPUT_COMMENT_RE = re.compile(r"<!--\s*legminus1 input:\s*(.+?)\s*-->")
_SOURCE_COMMENT_RE = re.compile(r"<!--\s*legminus1 source:\s*(.+?)\s*-->")
_LEAF_RE = re.compile(r"\{[$+*]?([A-Za-z_][\w.]*)\}")
_FINAL_RE = re.compile(r"(?m)^Final:\s*(.*)$")
_SCOPE_RE = re.compile(r"(?m)^-\s*Scope:\s*(.*)$")


# ---------------------------------------------------------------------------
# Input → plain text
# ---------------------------------------------------------------------------

def _doc_to_text(input_path: Path) -> str:
    """Convert a .docx/.pdf/.html(.htm) document to plain text for scanning."""
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        html = convert_docx(input_path)
    elif suffix == ".pdf":
        html = convert_pdf(input_path)
    elif suffix in (".html", ".htm"):
        html = input_path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"unsupported format '{suffix}'. Accepted: .docx, .pdf, .html")
    try:
        from bs4 import BeautifulSoup
        return BeautifulSoup(html, "html.parser").get_text()
    except ImportError:
        return re.sub(r"<[^>]+>", " ", html)


# ---------------------------------------------------------------------------
# Loop membership (mirrors leg0's [Name/]...[/Name] marker pairing)
# ---------------------------------------------------------------------------

def _loop_spans(text: str) -> list[tuple[int, int, str]]:
    """Return (inner_start, inner_end, name) for each matched [Name/]...[/Name]."""
    pending: dict[str, int] = {}
    spans: list[tuple[int, int, str]] = []
    for m in _LOOP_MARKER_RE.finditer(text):
        name = m.group(2)
        if m.group(1) != "/" and m.group(3) == "/":  # [Name/] opener
            pending[name] = m.end()
        elif m.group(1) == "/" and m.group(3) != "/" and name in pending:  # [/Name] closer
            spans.append((pending.pop(name), m.start(), name))
    return spans


def _loop_for_position(pos: int, spans: list[tuple[int, int, str]]) -> str | None:
    for start, end, name in spans:
        if start <= pos < end:
            return name
    return None


def collect_placeholders(text: str) -> list[dict]:
    """Extract unique ``{leaf}`` tokens with occurrence + loop membership.

    A leaf is treated as a loop field only when *every* occurrence sits inside
    the same loop (mirrors Leg 0's loop-field classification); otherwise it is
    document-level.
    """
    spans = _loop_spans(text)
    order: list[str] = []
    occ: dict[str, str] = {}
    loops_seen: dict[str, set] = {}
    for m in _FIELD_RE.finditer(text):
        symbol, name = m.group(1), m.group(2)
        if name not in occ:
            occ[name] = OCCURRENCE_SYMBOLS[symbol]
            order.append(name)
        loops_seen.setdefault(name, set()).add(_loop_for_position(m.start(), spans))

    fields: list[dict] = []
    for name in order:
        seen = loops_seen[name]
        # All occurrences inside one and the same loop → loop field.
        loop = next(iter(seen)) if len(seen) == 1 else None
        fields.append({"leaf": name, "occurrence": occ[name], "loop": loop})
    return fields


# ---------------------------------------------------------------------------
# Suggest mode
# ---------------------------------------------------------------------------

def resolve_fields(fields: list[dict], reg: dict, roots: list[str]) -> list[dict]:
    """Match each placeholder against the registry, attaching the verdict."""
    candidates = build_candidate_index(reg, roots=roots)
    iterables = {str(i.get("name") or "").lower() for i in reg.get("iterables") or []}
    results: list[dict] = []
    for f in fields:
        loop = f["loop"]
        warn = ""
        if loop and loop.lower() not in iterables:
            warn = f"loop `{loop}` is not a registry iterable — treating leaf as document-level"
            loop = None
        verdict = match_leaf(f["leaf"], loop, candidates)
        results.append({**f, "loop": loop, "warn": warn, **verdict})
    return results


def _scope_label(loop: str | None) -> str:
    return f"loop: {loop}" if loop else "document-level"


def write_path_review(results: list[dict], stem: str, source: str,
                      input_path: Path, roots: list[str], out: Path) -> None:
    roots_str = ", ".join(roots) if roots else "(none declared in filename)"
    lines = [
        f"# Path Review — {stem}",
        "",
        f"<!-- legminus1 input: {input_path.resolve()} -->",
        f"<!-- legminus1 source: {source} -->",
        "",
        "**System / canonical copy.** The customer fills the simpler "
        f"`{stem}.path-review.csv` (in `action-needed/`); its `final` column is "
        "written back onto the `Final:` lines below on apply. You can also edit the "
        "`Final:` lines here directly and re-run:",
        "",
        f"    python3 -m velocity_converter.legminus1_resolve_paths \\",
        f"        --parse-path-review {out.name} "
        f"--output-dir {machine_dir_for_action_file(out) or out.parent}",
        "",
        f"- Rendering root(s): {roots_str}",
        "- Accessors are **registry-matched only** — NOT verified against the "
        "compiled SDK (Leg 2 does that downstream).",
        "- Run `python3 -m velocity_converter.list_paths` to browse every accessor.",
        "",
    ]
    for r in results:
        lines += ["---", "", f"## Field: {{{r['leaf']}}}", ""]
        lines.append(f"- Scope: {_scope_label(r['loop'])}")
        lines.append(f"- Occurrence: {r['occurrence']}")
        if r["status"] == "resolved":
            lines.append(f"- Status: resolved ({r['match']})")
            lines.append(f"- Suggested: {r['chosen']}")
        elif r["status"] == "ambiguous":
            lines.append("- Status: AMBIGUOUS — pick one and put it on the Final line")
        else:
            lines.append("- Status: NO MATCH in registry — supply an accessor or fix the token")
        if r.get("scope_note"):
            lines.append(f"- Note: {r['scope_note']}")
        if r.get("warn"):
            lines.append(f"- Warning: {r['warn']}")
        if r["alternatives"]:
            lines.append("- Alternatives:")
            for a in r["alternatives"]:
                extra = ""
                if a.get("options"):
                    extra = f" — allowed: {', '.join(str(o) for o in a['options'][:6])}"
                if a.get("source") == "datafetcher":
                    extra += " — DataFetcher-sourced"
                lines.append(f"  - {a['accessor']} — {a['why']}{extra}")
        lines.append(f"Final: {r['chosen']}")
        lines.append("")
    out.write_text("\n".join(lines), encoding="utf-8")


def _candidate_accessors(r: dict) -> list[str]:
    """Accessor candidates for a leaf, top suggestion first, de-duplicated.

    The chosen/suggested accessor leads; the ranked alternatives follow. Used to
    build the multi-line ``suggested`` cell of the customer CSV.
    """
    accs: list[str] = []
    if r.get("chosen"):
        accs.append(r["chosen"])
    for a in r.get("alternatives") or []:
        acc = a.get("accessor")
        if acc and acc not in accs:
            accs.append(acc)
    return accs


def write_path_review_csv(results: list[dict], out: Path, *, append: bool = False,
                          no_suggest: bool = False) -> None:
    """Customer-facing view of the review — three columns: ``field`` /
    ``suggested`` / ``final``.

    ``suggested`` is a multi-line cell (one candidate accessor per line, the top
    pick first) so Excel stacks the options; ``final`` is the single accessor the
    customer keeps, pre-filled with the top pick (blank when nothing matched).
    The ``.path-review.md`` stays the canonical record — on apply the ``final``
    column is written back onto its ``Final:`` lines.

    ``no_suggest`` blanks both ``suggested`` and ``final`` so the customer maps
    every leaf by hand (the registry analysis still lands in the ``.md`` /
    ``.path-map.yaml`` — only this human-fill view is blanked).

    When ``append`` and ``out`` already exists, the existing data rows are kept
    verbatim (preserving any customer-edited ``final``) and only ``results`` whose
    ``{leaf}`` is not already a row are added. This is how a Pass-2 delta lands as
    *extra* rows on the full pass-1 review instead of replacing it, so the
    customer always sees every field in one file.
    """
    rows = [["field", "suggested", "final"]]
    existing_fields: set[str] = set()
    if append and out.is_file():
        with out.open(encoding="utf-8-sig", newline="") as fh:
            for i, row in enumerate(csv.reader(fh)):
                if i == 0 or not row:
                    continue  # skip header / blank lines
                rows.append(row)
                existing_fields.add((row[0] or "").strip())
    for r in results:
        field = f"{{{r['leaf']}}}"
        if field in existing_fields:
            continue  # already on the sheet (e.g. a field also used in the body)
        if no_suggest:
            rows.append([field, "", ""])
            continue
        cands = _candidate_accessors(r)
        suggested = "\n".join(cands) if cands else "(no registry match -- type an accessor)"
        final = cands[0] if cands else ""
        rows.append([field, suggested, final])
    with out.open("w", encoding="utf-8", newline="") as fh:
        # Default dialect: \r\n row terminator, fields with embedded newlines are
        # quoted — Excel shows the suggested cell as stacked lines.
        csv.writer(fh).writerows(rows)


def read_path_review_csv(csv_path: Path) -> dict[str, str]:
    """Parse a customer-filled ``.path-review.csv`` → ``{leaf: final accessor}``.

    ``final`` is meant to be a single accessor; if the customer left several
    candidate lines in the cell, the first non-empty line wins.
    """
    finals: dict[str, str] = {}
    with csv_path.open(encoding="utf-8-sig", newline="") as fh:
        for row in csv.DictReader(fh):
            m = _LEAF_RE.search((row.get("field") or "").strip())
            if not m:
                continue
            raw = (row.get("final") or "").strip()
            finals[m.group(1)] = next(
                (ln.strip() for ln in raw.splitlines() if ln.strip()), ""
            )
    return finals


def _patch_review_finals(md_path: Path, finals: dict[str, str]) -> None:
    """Write the CSV ``final`` values onto the canonical review's ``Final:``
    lines, block by block — keeps the ``.md`` the system source of truth."""
    text = md_path.read_text(encoding="utf-8")
    cur: str | None = None
    out_lines: list[str] = []
    for line in text.split("\n"):
        hm = re.match(r"##\s+Field:\s*\{[$+*]?([A-Za-z_][\w.]*)\}", line)
        if hm:
            cur = hm.group(1)
        if line.startswith("Final:") and cur in finals:
            out_lines.append(f"Final: {finals[cur]}")
        else:
            out_lines.append(line)
    md_path.write_text("\n".join(out_lines), encoding="utf-8")


def _datafetcher_for(accessor: str, results_alts: list[dict]) -> dict | None:
    for a in results_alts:
        if a["accessor"] == accessor and a.get("source") == "datafetcher":
            return a.get("datafetcher")
    return None


def write_path_map(results: list[dict], stem: str, source: str, input_path: Path,
                   roots: list[str], out: Path) -> None:
    fields = []
    for r in results:
        entry = {
            "leaf": r["leaf"],
            "occurrence": r["occurrence"],
            "scope": _scope_label(r["loop"]),
            "chosen": r["chosen"],
            "status": r["status"],
            "match": r["match"],
        }
        df = _datafetcher_for(r["chosen"], r["alternatives"]) if r["chosen"] else None
        if df:
            entry["datafetcher"] = df
        fields.append(entry)
    data = {
        "schema_version": "1.0",
        "generated_at": dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%S+00:00"),
        "source": source,
        "input_path": str(input_path.resolve()),
        "rendering_roots": roots,
        "fields": fields,
    }
    out.write_text(yaml.dump(data, default_flow_style=False, allow_unicode=True),
                   encoding="utf-8")


def write_path_changes(results: list[dict], stem: str, out: Path,
                        chosen_by: dict[str, str] | None = None) -> None:
    """Before/after audit. ``chosen_by`` maps leaf → 'legMinus1 (suggested)' |
    'human (override)' | 'human (selection)'; defaults to suggested/needs-human."""
    lines = [
        f"# Path Changes — {stem}",
        "",
        "Before/after record of every `{leaf}` the author wrote and the full "
        "accessor it resolved to. This is the traceability anchor — the original "
        "document is never modified.",
        "",
        "| Field (before) | Accessor (after) | Scope | Basis | Decided by |",
        "|---|---|---|---|---|",
    ]
    for r in results:
        after = r["chosen"] or "— (unresolved)"
        if r["status"] == "resolved":
            basis = "human-chosen" if r["match"] == "human" else f"{r['match']} match"
        elif r["status"] == "ambiguous":
            basis = f"{len(r['alternatives'])} candidates"
        else:
            basis = "no registry match"
        if chosen_by is not None:
            decided = chosen_by.get(r["leaf"], "—")
        else:
            decided = "legMinus1 (suggested)" if r["status"] == "resolved" else "needs human"
        lines.append(
            f"| `{{{r['leaf']}}}` | `{after}` | {_scope_label(r['loop'])} | {basis} | {decided} |"
        )
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _variants_csv_text(csv_path: Path) -> str:
    """Concatenate the ``text`` cells of a ``<stem>.variants.csv`` so
    :func:`collect_placeholders` can scan the ``{leaf}`` tokens the customer
    authored *in the CSV* (variants-only plan §2.6, Decision B).

    The variant text post-dates Leg -1 pass 1 (the source doc held only the
    ``[[$token]]`` marker), so its leaves were never seen by the first scan —
    feeding the CSV text in as a second source surfaces them.
    """
    from velocity_converter.condition_dsl import _read_csv_rows  # noqa: PLC0415
    try:
        rows = _read_csv_rows(csv_path.read_text(encoding="utf-8"))
    except OSError:
        return ""
    return "\n".join(r.get("text", "") for r in rows)


def run_suggest(
    input_path: Path,
    registry_path: Path | None,
    output_dir: Path,
    variants_csv: Path | None = None,
    no_suggest: bool = False,
) -> int:
    text = _doc_to_text(input_path)
    source = input_path.name
    roots, root_err = parse_rendering_roots(source)
    if root_err:
        print(f"WARNING: {root_err}", file=sys.stderr)
    # Registry is optional in no-suggest mode: with none loaded every leaf is
    # unmatched, which is fine — the CSV is blanked for manual fill regardless.
    reg = (yaml.safe_load(registry_path.read_text(encoding="utf-8")) or {}
           if registry_path and registry_path.exists() else {})

    # Pass 2 (Decision B): also scan the filled variants.csv's text cells.
    pass2 = variants_csv is not None and variants_csv.is_file()
    if pass2:
        text = text + "\n" + _variants_csv_text(variants_csv)

    fields = collect_placeholders(text)
    if not fields:
        print("No {leaf} placeholders found in the document.", file=sys.stderr)
        return 1
    results = resolve_fields(fields, reg, roots)

    stem = input_path.stem
    out_dir = output_dir / stem
    out_dir.mkdir(parents=True, exist_ok=True)
    # The customer fills the simple .path-review.csv (→ action-needed/); the
    # canonical .path-review.md and the machine map + audit stay in the stem's
    # output dir. On apply, the CSV's `final` column is written back onto the
    # md's `Final:` lines (the md remains the system source of truth).
    review = out_dir / f"{stem}.path-review.md"
    review_csv = action_needed_file(out_dir, f"{stem}.path-review.csv")
    pmap = out_dir / f"{stem}.path-map.yaml"
    changes = out_dir / f"{stem}.path-changes.md"

    if pass2:
        # Dedup against pass 1's path-map: leaves already resolved up front are
        # carried forward silently — only net-new (variant-text) leaves surface
        # into the delta review. Leave the pass-1 path-map/audit untouched; the
        # human edits the delta review and re-runs --parse-path-review, which
        # merges the net-new leaves into the existing map (run_apply unions the
        # prior path-map with the review).
        prior_leaves: set[str] = set()
        if pmap.exists():
            prior_data = yaml.safe_load(pmap.read_text(encoding="utf-8")) or {}
            prior_leaves = {f["leaf"] for f in (prior_data.get("fields") or [])}
        net_new = [r for r in results if r["leaf"] not in prior_leaves]
        if not net_new:
            print("Pass 2: no net-new variant-text leaves — path-review unchanged "
                  "(Decision B skip; nothing for the customer to resolve).")
            return 0
        write_path_review(net_new, stem, source, input_path, roots, review)
        # Append the net-new variant-text leaves onto the existing pass-1 CSV so
        # the customer-facing sheet keeps every plain field too (not just the
        # delta). The pass-1 path-map/audit stay untouched (merged on apply).
        write_path_review_csv(net_new, review_csv, append=True, no_suggest=no_suggest)
        n_res = sum(1 for r in net_new if r["status"] == "resolved")
        n_amb = sum(1 for r in net_new if r["status"] == "ambiguous")
        n_none = sum(1 for r in net_new if r["status"] == "unmatched")
        print(f"Wrote {review}")
        print(f"Wrote {review_csv}")
        print(f"\nPass 2: {len(net_new)} net-new variant-text leaf(s): "
              f"{n_res} resolved, {n_amb} ambiguous, {n_none} unmatched.")
        print(f"Fill the `final` column in {review_csv.name}, then re-run with "
              "--parse-path-review-csv (net-new leaves merge into the existing path-map).")
        return 0

    write_path_review(results, stem, source, input_path, roots, review)
    write_path_review_csv(results, review_csv, no_suggest=no_suggest)
    write_path_map(results, stem, source, input_path, roots, pmap)
    write_path_changes(results, stem, changes)

    n_res = sum(1 for r in results if r["status"] == "resolved")
    n_amb = sum(1 for r in results if r["status"] == "ambiguous")
    n_none = sum(1 for r in results if r["status"] == "unmatched")
    for p in (review_csv, review, pmap, changes):
        print(f"Wrote {p}")
    print(f"\n{len(results)} field(s): {n_res} resolved, {n_amb} ambiguous, {n_none} unmatched.")
    if n_amb or n_none:
        print(f"Fill the `final` column in {review_csv.name}, then re-run with "
              "--parse-path-review-csv.")
    return 0


# ---------------------------------------------------------------------------
# Apply mode (--parse-path-review)
# ---------------------------------------------------------------------------

def parse_path_review(md_path: Path) -> tuple[list[dict], Path | None, str]:
    """Parse a (possibly human-edited) review → (entries, input_path, source).

    Each entry: {leaf, scope, chosen}. ``chosen`` is the Final-line value. The
    file is split on ``## Field:`` headers so each block is parsed in isolation.
    """
    text = md_path.read_text(encoding="utf-8")
    im = _INPUT_COMMENT_RE.search(text)
    sm = _SOURCE_COMMENT_RE.search(text)
    input_path = Path(im.group(1)) if im else None
    source = sm.group(1) if sm else md_path.stem.replace(".path-review", "")
    entries: list[dict] = []
    for block in re.split(r"(?m)^##\s+Field:\s*", text)[1:]:
        lm = _LEAF_RE.match(block.strip())
        if not lm:
            continue
        fm = _FINAL_RE.search(block)
        scm = _SCOPE_RE.search(block)
        entries.append({
            "leaf": lm.group(1),
            "scope": (scm.group(1).strip() if scm else ""),
            "chosen": (fm.group(1).strip() if fm else ""),
        })
    return entries, input_path, source


def _sub_text(text: str, mapping: dict[str, str], replaced: set | None = None) -> str:
    """Rewrite ``{leaf}`` → ``{accessor}`` preserving any occurrence symbol."""
    def repl(m: re.Match) -> str:
        symbol, name = m.group(1), m.group(2)
        if name in mapping and mapping[name]:
            if replaced is not None:
                replaced.add(name)
            return "{" + symbol + mapping[name] + "}"
        return m.group(0)
    return _FIELD_RE.sub(repl, text)


def _rewrite_docx(src: Path, dst: Path, mapping: dict[str, str]) -> set:
    from docx import Document
    doc = Document(str(src))
    replaced: set = set()

    def fix_paragraph(p) -> None:
        for run in p.runs:
            new = _sub_text(run.text, mapping, replaced)
            if new != run.text:
                run.text = new
        joined = "".join(r.text for r in p.runs)
        # Placeholder split across runs → flatten into the first run.
        if any(n in mapping and mapping[n] for _, n in _FIELD_RE.findall(joined)):
            new = _sub_text(joined, mapping, replaced)
            if new != joined and p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""

    def walk_tables(tables) -> None:
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        fix_paragraph(p)
                    walk_tables(cell.tables)

    for p in doc.paragraphs:
        fix_paragraph(p)
    walk_tables(doc.tables)
    doc.save(str(dst))
    return replaced


def write_resolved_doc(input_path: Path, mapping: dict[str, str],
                       out_dir: Path, stem: str) -> tuple[Path | None, list[str]]:
    """Write a resolved copy of the source with full accessors baked in.

    Returns (path, missed_leaves). PDF input cannot be rewritten cleanly, so a
    resolved HTML is emitted instead with a warning.
    """
    suffix = input_path.suffix.lower()
    wanted = {leaf for leaf, acc in mapping.items() if acc}
    if suffix == ".docx":
        dst = out_dir / f"{stem}.resolved.docx"
        replaced = _rewrite_docx(input_path, dst, mapping)
        return dst, sorted(wanted - replaced)
    if suffix in (".html", ".htm"):
        dst = out_dir / f"{stem}.resolved{suffix}"
        original = input_path.read_text(encoding="utf-8")
        replaced: set = set()
        dst.write_text(_sub_text(original, mapping, replaced), encoding="utf-8")
        return dst, sorted(wanted - replaced)
    if suffix == ".pdf":
        dst = out_dir / f"{stem}.resolved.html"
        replaced: set = set()
        html = convert_pdf(input_path)
        dst.write_text(_sub_text(html, mapping, replaced), encoding="utf-8")
        print("WARNING: PDF input cannot be rewritten in place — wrote resolved HTML "
              f"({dst.name}) instead of a resolved PDF.", file=sys.stderr)
        return dst, sorted(wanted - replaced)
    return None, sorted(wanted)


def _scope_to_loop(scope: str) -> str | None:
    return scope.split(":", 1)[1].strip() if scope.startswith("loop:") else None


def run_apply(review_path: Path, output_dir: Path | None) -> int:
    entries, input_path, source = parse_path_review(review_path)
    # Machine artifacts (path-map, audit, resolved doc) belong in the stem's
    # output dir, not next to the human-fill review in action-needed/.
    out_dir = output_dir or machine_dir_for_action_file(review_path) or review_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = source.rsplit(".", 1)[0] if "." in source else source

    final_by_leaf = {e["leaf"]: e["chosen"] for e in entries}
    scope_by_leaf = {e["leaf"]: e["scope"] for e in entries}

    # The suggest-mode path-map carries the structured state (scope, occurrence,
    # original suggestion, status) — read it so the audit can show suggested-vs-
    # override provenance instead of flattening everything to "human".
    prior = out_dir / f"{stem}.path-map.yaml"
    prior_fields: list[dict] = []
    roots: list[str] = []
    if prior.exists():
        prior_data = yaml.safe_load(prior.read_text(encoding="utf-8")) or {}
        prior_fields = prior_data.get("fields") or []
        roots = prior_data.get("rendering_roots") or []
        if input_path is None and prior_data.get("input_path"):
            input_path = Path(prior_data["input_path"])
    prior_by_leaf = {f["leaf"]: f for f in prior_fields}

    # Iterate the prior field order, then append any net-new leaves the review
    # introduced (Decision B pass 2: variant-text leaves absent from pass 1's
    # path-map must merge in, not be dropped). Falls back to review order when
    # there is no prior map at all.
    leaves = [f["leaf"] for f in prior_fields]
    for leaf in final_by_leaf:
        if leaf not in prior_by_leaf:
            leaves.append(leaf)
    if not leaves:
        leaves = list(final_by_leaf)

    results: list[dict] = []
    decided: dict[str, str] = {}
    for leaf in leaves:
        orig = prior_by_leaf.get(leaf, {})
        final = final_by_leaf.get(leaf, orig.get("chosen", ""))
        scope = scope_by_leaf.get(leaf) or orig.get("scope", "document-level")
        loop = _scope_to_loop(scope)
        orig_status = orig.get("status", "")
        orig_chosen = orig.get("chosen", "")
        if not final:
            decided[leaf] = "needs human"
            match = ""
        elif orig_status == "resolved" and final == orig_chosen:
            decided[leaf] = "legMinus1 (suggested)"
            match = orig.get("match", "exact")
        elif orig_status == "resolved" and final != orig_chosen:
            decided[leaf] = "human (override)"
            match = "human"
        else:  # ambiguous / unmatched originally → human picked/supplied
            decided[leaf] = "human (selection)"
            match = "human"
        results.append({
            "leaf": leaf, "occurrence": orig.get("occurrence", "required"),
            "loop": loop, "chosen": final,
            "status": "resolved" if final else "unmatched",
            "match": match, "alternatives": [], "scope_note": "", "warn": "",
        })

    mapping = {r["leaf"]: r["chosen"] for r in results}
    unresolved = [r["leaf"] for r in results if not r["chosen"]]

    pmap = out_dir / f"{stem}.path-map.yaml"
    changes = out_dir / f"{stem}.path-changes.md"
    write_path_map(results, stem, source, input_path or Path(source), roots, pmap)
    write_path_changes(results, stem, changes, chosen_by=decided)
    print(f"Wrote {pmap}")
    print(f"Wrote {changes}")

    if unresolved:
        print(f"WARNING: {len(unresolved)} field(s) still have an empty Final line: "
              f"{', '.join(unresolved)} — they will be left as-is for Leg 0.",
              file=sys.stderr)

    if input_path and input_path.exists():
        resolved, missed = write_resolved_doc(input_path, mapping, out_dir, stem)
        if resolved:
            print(f"Wrote {resolved}")
        if missed:
            print(f"WARNING: {len(missed)} placeholder(s) not found in the source "
                  f"document (not substituted): {', '.join(missed)}", file=sys.stderr)
    else:
        print("Note: original document not found — wrote map + audit only "
              "(leg0 can still consume the path-map).", file=sys.stderr)
    return 0


def run_apply_csv(csv_path: Path, output_dir: Path | None) -> int:
    """Apply a customer-filled ``.path-review.csv``: write its ``final`` column
    onto the canonical ``.path-review.md``, then run the normal md apply.

    The md is the system source of truth; the CSV is just the customer surface,
    so we fold the CSV back into the md and reuse :func:`run_apply` unchanged.
    """
    out_dir = output_dir or machine_dir_for_action_file(csv_path) or csv_path.parent
    suffix = ".path-review.csv"
    stem = (csv_path.name[: -len(suffix)] if csv_path.name.endswith(suffix)
            else csv_path.stem.replace(".path-review", ""))
    md_path = out_dir / f"{stem}.path-review.md"
    if not md_path.exists():
        print(f"Error: canonical review not found next to the CSV: {md_path}\n"
              "Run Leg -1 suggest first — it writes the .path-review.md the CSV maps onto.",
              file=sys.stderr)
        return 1
    _patch_review_finals(md_path, read_path_review_csv(csv_path))
    print(f"Folded {csv_path.name} → {md_path}")
    return run_apply(md_path, out_dir)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _find_registry(start: Path) -> Path | None:
    cur = start.resolve()
    for _ in range(8):
        for rel in ("registry/path-registry.yaml", "path-registry.yaml"):
            cand = cur / rel
            if cand.is_file():
                return cand
        if cur.parent == cur:
            break
        cur = cur.parent
    return None


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Leg -1: resolve bare {leaf} placeholders to full accessor paths."
    )
    parser.add_argument("--input", default=None, help="Path to .docx/.pdf/.html document")
    parser.add_argument("--registry", default=None, help="Path to path-registry.yaml")
    parser.add_argument("--output-dir", default=None, help="Output directory")
    parser.add_argument("--parse-path-review", default=None, metavar="REVIEW.md",
                        help="Parse a (human-edited) review → final map + resolved doc")
    parser.add_argument("--parse-path-review-csv", default=None, metavar="REVIEW.csv",
                        help="Apply a customer-filled .path-review.csv: fold its `final` "
                             "column onto the canonical .path-review.md, then run the md apply")
    parser.add_argument("--variants-csv", default=None, metavar="VARIANTS.csv",
                        help="Pass 2 (Decision B): also scan a filled <stem>.variants.csv's "
                             "text cells; emit a delta review of net-new variant-text leaves "
                             "(deduped against the pass-1 path-map)")
    parser.add_argument("--no-suggest", action="store_true",
                        help="Blank the `suggested`/`final` columns in the path-review.csv "
                             "so the customer maps every leaf by hand (the .md/.path-map "
                             "still carry the registry analysis)")
    args = parser.parse_args()

    if args.parse_path_review_csv:
        review = Path(args.parse_path_review_csv)
        if not review.exists():
            print(f"Error: file not found: {review}", file=sys.stderr)
            return 1
        out_dir = Path(args.output_dir) if args.output_dir else None
        return run_apply_csv(review, out_dir)

    if args.parse_path_review:
        review = Path(args.parse_path_review)
        if not review.exists():
            print(f"Error: file not found: {review}", file=sys.stderr)
            return 1
        out_dir = Path(args.output_dir) if args.output_dir else None
        return run_apply(review, out_dir)

    if not args.input:
        print("Error: --input is required (unless using --parse-path-review[-csv])", file=sys.stderr)
        return 1
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        return 1

    registry_path = Path(args.registry) if args.registry else _find_registry(input_path.parent)
    if not registry_path or not registry_path.exists():
        if args.no_suggest:
            # No-suggest can run without a registry — emit blank fill files anyway.
            print("No registry loaded — emitting blank path-review.csv for manual fill.",
                  file=sys.stderr)
            registry_path = None
        else:
            print("Error: registry not found — pass --registry registry/path-registry.yaml",
                  file=sys.stderr)
            return 1

    output_dir = Path(args.output_dir) if args.output_dir else input_path.parent
    variants_csv = Path(args.variants_csv) if args.variants_csv else None
    if variants_csv and not variants_csv.exists():
        print(f"Error: variants CSV not found: {variants_csv}", file=sys.stderr)
        return 1
    return run_suggest(input_path, registry_path, output_dir, variants_csv=variants_csv,
                       no_suggest=args.no_suggest)


if __name__ == "__main__":
    sys.exit(main())
